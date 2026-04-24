import os
import streamlit as st
import requests
import json
import re
from docx import Document
from docx.shared import Pt
from io import BytesIO
import pandas as pd

BACKEND         = os.getenv("BACKEND_URL", "http://127.0.0.1:9000")
CITERAG_BACKEND = BACKEND + "/rag"  # CiteRAG routes now on same backend

st.set_page_config(page_title="DocWizard", page_icon="🪄", layout="wide", initial_sidebar_state="expanded")

# ─────────────────────────────────────────────────────────────────────────────
# CSS — Navy + Gold theme
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
.stApp { background: linear-gradient(135deg, #0a0e1a 0%, #0d1425 50%, #0a0e1a 100%); }
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #0d1631 0%, #0a1128 100%) !important;
    border-right: 1px solid #1e2d4a !important;
}
[data-testid="stSidebar"] .stRadio label { color: #a0b4cc !important; font-size: 0.92rem; padding: 6px 0; }
[data-testid="stSidebar"] .stRadio label:hover { color: #f5c842 !important; }
.page-header {
    background: linear-gradient(90deg, #0d1631, #112044);
    border: 1px solid #1e3a6e; border-left: 4px solid #f5c842;
    border-radius: 12px; padding: 24px 32px; margin-bottom: 28px;
}
.page-header h1 { font-size: 1.9rem; font-weight: 700; color: #ffffff; margin: 0 0 4px 0; }
.page-header p { color: #7a94b8; font-size: 0.88rem; margin: 0; }
.step-card {
    background: linear-gradient(135deg, #0f1a2e 0%, #111f38 100%);
    border: 1px solid #1e3055; border-radius: 14px; padding: 24px 28px; margin-bottom: 20px;
}
.step-label { font-size: 0.7rem; font-weight: 700; letter-spacing: 2px; color: #f5c842; text-transform: uppercase; margin-bottom: 4px; }
.doc-card {
    background: linear-gradient(135deg, #0f1a2e 0%, #0d1731 100%);
    border: 1px solid #1e3055; border-radius: 14px; padding: 32px 36px;
    margin-bottom: 24px; box-shadow: 0 4px 24px rgba(0,0,0,0.3);
}
.doc-title { font-size: 1.7rem; font-weight: 700; color: #ffffff; border-bottom: 2px solid #f5c842; padding-bottom: 12px; margin-bottom: 24px; }
.doc-section-heading { font-size: 1.05rem; font-weight: 600; color: #f5c842; margin-top: 24px; margin-bottom: 8px; padding-left: 12px; border-left: 3px solid #f5c842; }
.doc-paragraph { font-size: 0.93rem; color: #c8d8ec; line-height: 1.85; margin-bottom: 10px; padding-left: 15px; }
.bullet-item { color: #c8d8ec; font-size: 0.93rem; line-height: 1.85; padding-left: 22px; position: relative; margin-bottom: 4px; }
.bullet-item::before { content: "▸"; color: #f5c842; position: absolute; left: 6px; font-size: 0.8rem; }
.format-badge-excel { background: rgba(245,200,66,0.12); color: #f5c842; border: 1px solid #f5c842; border-radius: 20px; padding: 3px 12px; font-size: 11px; font-weight: 700; letter-spacing: 0.5px; }
.format-badge-word { background: rgba(59,130,246,0.12); color: #60a5fa; border: 1px solid #3b82f6; border-radius: 20px; padding: 3px 12px; font-size: 11px; font-weight: 700; letter-spacing: 0.5px; }
.hist-card { background: linear-gradient(135deg, #0f1a2e 0%, #111f38 100%); border: 1px solid #1e3055; border-radius: 12px; padding: 18px 22px; margin-bottom: 14px; transition: border-color 0.2s; }
.hist-card:hover { border-color: #f5c842; }
.hist-title { font-size: 1rem; font-weight: 700; color: #e8f0ff; }
.hist-meta { font-size: 0.8rem; color: #6b82a8; margin-top: 5px; }
.stButton > button {
    background: linear-gradient(135deg, #1a3a6e, #1e4494) !important;
    color: #ffffff !important; border: 1px solid #2a52a8 !important;
    border-radius: 8px !important; font-weight: 600 !important;
    font-size: 0.88rem !important; padding: 0.45rem 1.2rem !important; transition: all 0.2s !important;
}
.stButton > button:hover {
    background: linear-gradient(135deg, #f5c842, #e6b830) !important;
    color: #0a0e1a !important; border-color: #f5c842 !important;
    transform: translateY(-1px); box-shadow: 0 4px 12px rgba(245,200,66,0.3) !important;
}
.stButton > button:disabled {
    background: #0f1a2e !important;
    color: #2a3a5a !important;
    border: 1px solid #1a2a44 !important;
    cursor: not-allowed !important;
    opacity: 0.45 !important;
    transform: none !important;
    box-shadow: none !important;
}
.stDownloadButton > button { background: linear-gradient(135deg, #f5c842, #e6b830) !important; color: #0a0e1a !important; border: none !important; border-radius: 8px !important; font-weight: 700 !important; font-size: 0.9rem !important; }
.stDownloadButton > button:hover { background: linear-gradient(135deg, #ffe066, #f5c842) !important; box-shadow: 0 4px 16px rgba(245,200,66,0.4) !important; }
.stTextInput > div > div > input, .stTextArea > div > div > textarea { background: #0d1631 !important; border: 1px solid #1e3055 !important; border-radius: 8px !important; color: #e8f0ff !important; font-size: 0.92rem !important; }
.stTextInput > div > div > input:focus, .stTextArea > div > div > textarea:focus { border-color: #f5c842 !important; box-shadow: 0 0 0 2px rgba(245,200,66,0.15) !important; }
.streamlit-expanderHeader { background: #0f1a2e !important; border: 1px solid #1e3055 !important; border-radius: 8px !important; color: #c8d8ec !important; font-weight: 600 !important; }
.streamlit-expanderContent { background: #0d1631 !important; border: 1px solid #1e3055 !important; border-top: none !important; }
.stProgress > div > div > div > div { background: linear-gradient(90deg, #1e4494, #f5c842) !important; border-radius: 4px !important; }
.stSuccess { background: rgba(245,200,66,0.08) !important; border: 1px solid rgba(245,200,66,0.3) !important; border-radius: 8px !important; color: #f5c842 !important; }
.stInfo { background: rgba(59,130,246,0.08) !important; border: 1px solid rgba(59,130,246,0.3) !important; border-radius: 8px !important; }
hr { border-color: #1e3055 !important; margin: 20px 0 !important; }
[data-testid="stDataFrame"] { border: 1px solid #1e3055 !important; border-radius: 10px !important; overflow: hidden; }
.stSubheader, h2, h3 { color: #e8f0ff !important; }
.sidebar-logo { text-align: center; padding: 16px 0 8px 0; }
.sidebar-logo-text { font-size: 1.3rem; font-weight: 800; background: linear-gradient(90deg, #f5c842, #60a5fa); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text; letter-spacing: -0.5px; }
.sidebar-logo-sub { font-size: 0.72rem; color: #4a6080; letter-spacing: 2px; text-transform: uppercase; }
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# Helpers — IST time conversion
# ─────────────────────────────────────────────────────────────────────────────

def to_ist(dt_str: str) -> str:
    if not dt_str:
        return ""
    try:
        from datetime import datetime, timezone, timedelta
        IST = timezone(timedelta(hours=5, minutes=30))
        dt_str_clean = dt_str.replace("T", " ").split(".")[0]
        dt = datetime.strptime(dt_str_clean, "%Y-%m-%d %H:%M:%S")
        dt_utc = dt.replace(tzinfo=timezone.utc)
        return dt_utc.astimezone(IST).strftime("%d %b %Y, %I:%M %p")
    except Exception:
        return dt_str


# ─────────────────────────────────────────────────────────────────────────────
# Notion helpers — smart push/update dialog
# ─────────────────────────────────────────────────────────────────────────────

def _do_notion_push(notion_title: str, doc_format: str, doc_content: dict,
                    db_id: int, notion_key: str, version: int = 1, doc_type: str = "General"):
    """Execute a fresh Notion push and show result."""
    with st.spinner("Pushing to Notion..."):
        res = requests.post(f"{BACKEND}/notion/push", json={
            "title":      notion_title,
            "doc_format": doc_format,
            "content":    doc_content,
            "db_id":      db_id,
            "version":    version,
            "doc_type":   doc_type,
        })
    if res.status_code == 200:
        data = res.json()
        st.session_state[notion_key] = data.get("page_id")
        notion_url = data.get("url", "")
        st.success("✅ Pushed to Notion!")
        st.markdown(
            f'<a href="{notion_url}" target="_blank" style="display:inline-block;margin-top:8px;'
            f'padding:8px 18px;background:linear-gradient(135deg,#f5c842,#e6b830);color:#0a0e1a;'
            f'font-weight:700;border-radius:8px;text-decoration:none;font-size:0.88rem;">'
            f'🔗 Open in Notion</a>',
            unsafe_allow_html=True
        )
    else:
        st.error(f"Notion push failed: {res.text[:200]}")


def _do_notion_update(existing_page_id: str, notion_title: str, doc_format: str,
                      doc_content: dict, notion_key: str, version: int = 1, doc_type: str = "General"):
    """Execute a Notion page update and show result."""
    with st.spinner("Updating Notion page..."):
        res = requests.post(f"{BACKEND}/notion/update", json={
            "page_id":    existing_page_id,
            "title":      notion_title,
            "doc_format": doc_format,
            "content":    doc_content,
            "version":    version,
            "doc_type":   doc_type,
        })
    if res.status_code == 200:
        notion_url = res.json().get("url", "")
        st.success("✅ Notion page updated!")
        st.markdown(
            f'<a href="{notion_url}" target="_blank" style="display:inline-block;margin-top:8px;'
            f'padding:8px 18px;background:linear-gradient(135deg,#f5c842,#e6b830);color:#0a0e1a;'
            f'font-weight:700;border-radius:8px;text-decoration:none;font-size:0.88rem;">'
            f'🔗 Open in Notion</a>',
            unsafe_allow_html=True
        )
    elif res.status_code == 500 and "404" in res.text:
        st.warning("Previous Notion page no longer exists — creating a new one...")
        st.session_state.pop(notion_key, None)
    else:
        st.error(f"Update failed: {res.text[:200]}")


def _resolve_notion_version(title: str, db_id: int | None) -> tuple[str, int]:
    """Return (notion_title, version) by reading the doc's version from DB."""
    version = 1
    notion_title = title
    if db_id:
        ver_check = requests.get(f"{BACKEND}/documents/{db_id}/versions")
        if ver_check.status_code == 200:
            versions = ver_check.json().get("versions", [])
            v_num = next((v["version"] for v in versions if v["id"] == db_id), None)
            if v_num:
                version = v_num
                if v_num > 1:
                    notion_title = f"{title} (v{v_num})"
    return notion_title, version


def push_to_notion_ui(title: str, doc_format: str, doc_content: dict,
                      db_id: int = None, notion_key: str = "notion_page_id",
                      doc_type: str = None):
    """
    Trigger a Notion push. Stores intent in session state so the
    dialog renders at the top level on the next rerun — not inside
    a button callback where it would vanish immediately.
    """
    pending_key = f"notion_pending_{notion_key}"
    existing_page_id = st.session_state.get(notion_key)

    # Use passed doc_type, fallback to session state (generate page), then "General"
    doc_type = doc_type or st.session_state.get("doc_type", "General") or "General"

    if not existing_page_id:
        # First push — store intent and rerun so it renders at top level
        st.session_state[pending_key] = {
            "action": "push", "title": title, "doc_format": doc_format,
            "doc_content": doc_content, "db_id": db_id, "notion_key": notion_key,
            "doc_type": doc_type,
        }
    else:
        # Already pushed — store "ask" intent, dialog shown at top level
        st.session_state[pending_key] = {
            "action": "ask", "title": title, "doc_format": doc_format,
            "doc_content": doc_content, "db_id": db_id, "notion_key": notion_key,
            "existing_page_id": existing_page_id, "doc_type": doc_type,
        }
    st.rerun()


def _render_notion_dialogs():
    """
    Called once at the TOP of each page render.
    Processes any pending Notion actions stored in session state.
    """
    for key in list(st.session_state.keys()):
        if not key.startswith("notion_pending_"):
            continue
        p = st.session_state[key]
        notion_key   = p["notion_key"]
        title        = p["title"]
        doc_format   = p["doc_format"]
        doc_content  = p["doc_content"]
        db_id        = p.get("db_id")
        doc_type     = p.get("doc_type", "General")
        notion_title, version = _resolve_notion_version(title, db_id)

        if p["action"] == "push":
            st.session_state.pop(key)
            _do_notion_push(notion_title, doc_format, doc_content, db_id, notion_key, version, doc_type)

        elif p["action"] == "ask":
            existing_page_id = p["existing_page_id"]
            st.info(f"**'{title}'** was already pushed to Notion. What would you like to do?")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("🔄 Update existing page", key=f"nd_upd_{notion_key}", use_container_width=True):
                    st.session_state.pop(key)
                    _do_notion_update(existing_page_id, notion_title, doc_format,
                                      doc_content, notion_key, version, doc_type)
                    st.rerun()
            with col2:
                if st.button("➕ Create new page", key=f"nd_new_{notion_key}", use_container_width=True):
                    st.session_state.pop(key)
                    st.session_state.pop(notion_key, None)
                    _do_notion_push(notion_title, doc_format, doc_content, db_id, notion_key, version, doc_type)
                    st.rerun()

        elif p["action"] == "update":
            existing_page_id = p["existing_page_id"]
            st.session_state.pop(key)
            _do_notion_update(existing_page_id, notion_title, doc_format,
                              doc_content, notion_key, version, doc_type)

        elif p["action"] == "new":
            st.session_state.pop(key)
            st.session_state.pop(notion_key, None)
            _do_notion_push(notion_title, doc_format, doc_content, db_id, notion_key, version, doc_type)


# ─────────────────────────────────────────────────────────────────────────────
# Shared doc card renderer — used by history page AND generate page
# ─────────────────────────────────────────────────────────────────────────────

def _render_doc_card(doc: dict, in_memory_content=None, in_memory_buf: "BytesIO|None" = None):
    """
    Render a single document card with all 6 actions.

    doc keys (required):
        title, doc_format, doc_id (int or None), version, doc_type,
        file_ext, created_at (str, may be empty)

    in_memory_content : dict — pass generated_sections / excel_data when the
                                doc is not yet saved (doc_id is None).
    in_memory_buf     : BytesIO — compiled file bytes for download when unsaved.
    """
    doc_id     = doc.get("doc_id")          # None = unsaved in-memory doc
    title      = doc["title"]
    doc_format = doc.get("doc_format", "word")
    file_ext   = doc.get("file_ext") or ("xlsx" if doc_format == "excel" else "docx")
    version    = doc.get("version", 1)
    created_at = to_ist(doc.get("created_at", ""))
    badge_cls  = "format-badge-excel" if doc_format == "excel" else "format-badge-word"
    badge_lbl  = "📊 Excel" if doc_format == "excel" else "📄 Word"
    card_key   = doc_id if doc_id is not None else f"mem_{title[:20]}"
    is_saved   = doc_id is not None

    # ── Card header ───────────────────────────────────────────────────────────
    ts_part = f"&nbsp;|&nbsp; 🕒 {created_at}" if created_at else ""
    id_part = f"&nbsp;|&nbsp; ID #{doc_id}"    if is_saved   else "&nbsp;|&nbsp; <em>not saved to DB yet</em>"
    _hdr_col, _rename_btn_col = st.columns([10, 1])
    with _hdr_col:
        st.markdown(f"""
        <div class="hist-card">
            <span class="hist-title">{title}</span>
            &nbsp;&nbsp;<span class="{badge_cls}">{badge_lbl}</span>
            &nbsp;&nbsp;<span style="background:rgba(99,179,237,0.15);color:#63b3ed;
                border:1px solid rgba(99,179,237,0.35);border-radius:4px;
                padding:2px 8px;font-size:0.72rem;font-weight:700;">v{version}</span>
            <div class="hist-meta">{ts_part}{id_part}</div>
        </div>
        """, unsafe_allow_html=True)
    with _rename_btn_col:
        st.markdown("<div style='margin-top:10px'></div>", unsafe_allow_html=True)
        if st.button("✏️", key=f"rename_trigger_{card_key}", help="Rename document"):
            st.session_state[f"renaming_{card_key}"] = not st.session_state.get(f"renaming_{card_key}", False)
            st.rerun()

    # ── Rename title ──────────────────────────────────────────────────────────
    _rename_key = f"renaming_{card_key}"
    if st.session_state.get(_rename_key):
        _rn_col1, _rn_col2, _rn_col3 = st.columns([6, 1, 1])
        with _rn_col1:
            _new_name = st.text_input(
                "New title", value=title, key=f"rename_input_{card_key}",
                label_visibility="collapsed"
            )
        with _rn_col2:
            if st.button("✅", key=f"rename_save_{card_key}", help="Save"):
                _new_name = _new_name.strip()
                if _new_name and _new_name != title:
                    # Update session state if this is the active generate doc
                    if st.session_state.get("title") == title:
                        st.session_state["title"] = _new_name
                        if "generated_sections" in st.session_state:
                            st.session_state["generated_sections"]["__title__"] = _new_name
                    # Update DB if saved
                    if is_saved:
                        _upd = requests.patch(f"{BACKEND}/documents/{doc_id}/rename",
                                              json={"title": _new_name})
                        if _upd.status_code == 200:
                            st.success(f"Renamed to '{_new_name}'")
                        else:
                            st.error("Rename failed")
                st.session_state[_rename_key] = False
                st.rerun()
        with _rn_col3:
            if st.button("✕", key=f"rename_cancel_{card_key}", help="Cancel"):
                st.session_state[_rename_key] = False
                st.rerun()

    col1, col2, col3, col4, col5, col6 = st.columns([2, 2, 2, 2, 2, 2])

    # ── 1. View ───────────────────────────────────────────────────────────────
    with col1:
        if st.button("👁 View", key=f"btn_view_{card_key}", use_container_width=True):
            st.session_state[f"view_{card_key}"] = not st.session_state.get(f"view_{card_key}", False)

    # ── 2. Versions (only for saved docs) ─────────────────────────────────────
    with col2:
        if is_saved:
            if st.button("🕓 Versions", key=f"btn_ver_{card_key}", use_container_width=True):
                st.session_state[f"ver_{card_key}"] = not st.session_state.get(f"ver_{card_key}", False)
        else:
            st.button("🕓 Versions", key=f"btn_ver_{card_key}", use_container_width=True, disabled=True, help="Save to DB first to view versions")

    # ── 3. Download ───────────────────────────────────────────────────────────
    with col3:
        mime = (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            if file_ext == "xlsx"
            else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        safe_fname = title.replace(" ", "_")[:40]

        if in_memory_buf is not None:
            # In-memory: direct download button
            in_memory_buf.seek(0)
            st.download_button(
                label=f"⬇️ Download .{file_ext}",
                data=in_memory_buf.read(),
                file_name=f"{safe_fname}.{file_ext}",
                mime=mime,
                use_container_width=True,
                key=f"dl_mem_{card_key}",
            )
        elif is_saved:
            if st.button("⬇️ Download", use_container_width=True, key=f"dl_trigger_{card_key}"):
                st.session_state[f"dl_ready_{card_key}"] = True
            if st.session_state.get(f"dl_ready_{card_key}"):
                dl_res = requests.get(f"{BACKEND}/documents/{doc_id}/download")
                if dl_res.status_code == 200:
                    st.download_button(
                        label=f"⬇️ Save .{file_ext}",
                        data=dl_res.content,
                        file_name=f"{safe_fname}_v{version}.{file_ext}",
                        mime=mime,
                        use_container_width=True,
                        key=f"dl_card_{card_key}",
                    )
                else:
                    st.caption("File not available.")
                    st.session_state.pop(f"dl_ready_{card_key}", None)
        else:
            st.button("⬇️ Download", key=f"dl_trigger_{card_key}", use_container_width=True, disabled=True)

    # ── 4. Push to Notion ─────────────────────────────────────────────────────
    with col4:
        _notion_key   = f"notion_hist_{card_key}"
        _notion_url   = st.session_state.get(f"notion_url_{card_key}")
        _notion_err   = st.session_state.get(f"notion_err_{card_key}")
        _already      = st.session_state.get(_notion_key)
        _nlbl         = "🔄 Update Notion" if _already else "🔗 Push to Notion"
        _doc_type_card = doc.get("doc_type", "General") or "General"

        if st.button(_nlbl, key=f"notion_card_{card_key}", use_container_width=True):
            # Resolve content
            _nc = in_memory_content
            _nfmt = doc_format
            _ntype = _doc_type_card
            _nver  = doc.get("version", 1)
            _ndb   = doc_id

            if _nc is None and is_saved:
                with st.spinner("Loading…"):
                    _det = requests.get(f"{BACKEND}/documents/{doc_id}")
                if _det.status_code == 200:
                    _dj   = _det.json()
                    _nc   = _dj.get("content", {})
                    _nfmt = _dj.get("doc_format", doc_format)
                    _ntype = _dj.get("doc_type", _doc_type_card) or _doc_type_card
                    _nver  = _dj.get("version", 1)
                else:
                    st.session_state[f"notion_err_{card_key}"] = "Could not load document content."
                    st.rerun()

            if _nc is not None:
                if not _already:
                    # Fresh push
                    with st.spinner("Pushing to Notion…"):
                        _nr = requests.post(f"{BACKEND}/notion/push", json={
                            "title": title, "doc_format": _nfmt,
                            "content": _nc, "db_id": _ndb,
                            "version": _nver, "doc_type": _ntype,
                        })
                    if _nr.status_code == 200:
                        _nd = _nr.json()
                        st.session_state[_notion_key]             = _nd.get("page_id")
                        st.session_state[f"notion_url_{card_key}"] = _nd.get("url", "")
                        st.session_state.pop(f"notion_err_{card_key}", None)
                    else:
                        st.session_state[f"notion_err_{card_key}"] = f"Push failed: {_nr.text[:200]}"
                else:
                    # Already pushed — show update/new dialog via session flag
                    st.session_state[f"notion_ask_{card_key}"] = {
                        "content": _nc, "fmt": _nfmt, "type": _ntype,
                        "ver": _nver, "db": _ndb, "page_id": _already,
                    }
                st.rerun()

        # Persistent result display
        if _notion_url:
            st.markdown(
                f'<a href="{_notion_url}" target="_blank" style="display:block;text-align:center;'
                f'margin-top:4px;padding:4px 8px;background:linear-gradient(135deg,#f5c842,#e6b830);'
                f'color:#0a0e1a;font-weight:700;border-radius:6px;text-decoration:none;font-size:0.78rem;">'
                f'🔗 Open in Notion</a>',
                unsafe_allow_html=True
            )
        if _notion_err:
            st.error(_notion_err)

    # ── 4b. Notion update/new dialog — full width below row ───────────────────
    _nask = st.session_state.get(f"notion_ask_{card_key}")
    if _nask:
        st.info(f"**'{title}'** was already pushed to Notion. What would you like to do?")
        _na1, _na2, _na3 = st.columns([2, 2, 1])
        with _na1:
            if st.button("🔄 Update existing page", key=f"nd_upd_{card_key}", use_container_width=True, type="primary"):
                with st.spinner("Updating Notion…"):
                    _ur = requests.post(f"{BACKEND}/notion/update", json={
                        "page_id":    _nask["page_id"],
                        "title":      title,
                        "doc_format": _nask["fmt"],
                        "content":    _nask["content"],
                        "version":    _nask["ver"],
                        "doc_type":   _nask["type"],
                    })
                if _ur.status_code == 200:
                    _ud = _ur.json()
                    st.session_state[f"notion_url_{card_key}"] = _ud.get("url", "")
                    st.session_state.pop(f"notion_err_{card_key}", None)
                else:
                    st.session_state[f"notion_err_{card_key}"] = f"Update failed: {_ur.text[:200]}"
                st.session_state.pop(f"notion_ask_{card_key}", None)
                st.rerun()
        with _na2:
            if st.button("➕ Create new page", key=f"nd_new_{card_key}", use_container_width=True):
                with st.spinner("Pushing to Notion…"):
                    _pr = requests.post(f"{BACKEND}/notion/push", json={
                        "title": title, "doc_format": _nask["fmt"],
                        "content": _nask["content"], "db_id": _nask["db"],
                        "version": _nask["ver"], "doc_type": _nask["type"],
                    })
                if _pr.status_code == 200:
                    _pd = _pr.json()
                    st.session_state[_notion_key]             = _pd.get("page_id")
                    st.session_state[f"notion_url_{card_key}"] = _pd.get("url", "")
                    st.session_state.pop(f"notion_err_{card_key}", None)
                else:
                    st.session_state[f"notion_err_{card_key}"] = f"Push failed: {_pr.text[:200]}"
                st.session_state.pop(f"notion_ask_{card_key}", None)
                st.rerun()
        with _na3:
            if st.button("✕ Cancel", key=f"nd_cancel_{card_key}", use_container_width=True):
                st.session_state.pop(f"notion_ask_{card_key}", None)
                st.rerun()

    # ── 5. Save to DB ─────────────────────────────────────────────────────────
    with col5:
        _saved_db = st.session_state.get(f"db_saved_{card_key}", False)
        if _saved_db or is_saved:
            lbl = "✅ In DB" if is_saved and not _saved_db else "✅ Saved"
            st.button(lbl, key=f"save_db_{card_key}", use_container_width=True, disabled=True)
        else:
            if st.button("💾 Save to DB", key=f"save_db_{card_key}", use_container_width=True, type="primary"):
                st.session_state[f"show_db_dialog_{card_key}"] = True

    # ── 6. Delete (only for saved docs) ───────────────────────────────────────
    with col6:
        if is_saved:
            if st.button("🗑 Delete", key=f"btn_del_{card_key}", use_container_width=True):
                del_res = requests.delete(f"{BACKEND}/documents/{doc_id}")
                if del_res.status_code == 200:
                    st.success(f"Deleted '{title}'")
                    st.rerun()
                else:
                    st.error("Delete failed")
        else:
            st.button("🗑 Delete", key=f"btn_del_{card_key}", use_container_width=True, disabled=True, help="Save to DB first to delete")

    # ── Save to DB dialog — full-width below button row ───────────────────────
    if st.session_state.get(f"show_db_dialog_{card_key}"):
        # Resolve content: prefer in-memory, else fetch from API
        _content_to_save = in_memory_content
        _type_to_save    = doc.get("doc_type", "General")
        _fmt_to_save     = doc_format
        _ext_to_save     = file_ext  # already None-safe from card extraction

        if _content_to_save is None and is_saved:
            _det2 = requests.get(f"{BACKEND}/documents/{doc_id}")
            if _det2.status_code == 200:
                _d2 = _det2.json()
                _content_to_save = _d2.get("content", {})
                _type_to_save    = _d2.get("doc_type", "General")
                _fmt_to_save     = _d2.get("doc_format", doc_format)
                _d2_ext          = _d2.get("file_ext")
                _ext_to_save     = _d2_ext or ("xlsx" if _fmt_to_save == "excel" else "docx")
            else:
                st.error("Could not load document details.")
                st.session_state.pop(f"show_db_dialog_{card_key}", None)
                return

        st.markdown(
            f'<div style="background:#0d1836;border:1px solid #1e3055;'            f'border-left:3px solid #f5c842;border-radius:8px;'            f'padding:14px 18px;margin:6px 0 10px;">'            f'<span style="font-size:0.9rem;font-weight:700;color:#fff;">💾 Save to DB</span>'            f'&nbsp;&nbsp;<span style="font-size:0.78rem;color:#7a94b8;">{title}</span>'            f'</div>',
            unsafe_allow_html=True
        )
        _chk = requests.post(f"{BACKEND}/documents/check-version", json={"title": title})
        _exists     = _chk.json().get("exists", False)     if _chk.status_code == 200 else False
        _latest_ver = _chk.json().get("latest_version", 1) if _chk.status_code == 200 else 1
        _latest_id  = _chk.json().get("latest_id")         if _chk.status_code == 200 else None

        if not _exists:
            _sv = requests.post(f"{BACKEND}/documents/save", json={
                "title": title, "doc_type": _type_to_save,
                "doc_format": _fmt_to_save, "content": _content_to_save,
                "file_ext": _ext_to_save, "save_mode": "new_version",
            })
            if _sv.status_code == 200:
                _svd = _sv.json()
                st.success(f"✅ Saved **'{title}'** as v{_svd.get('version',1)} (ID #{_svd.get('id')})")
                st.session_state[f"db_saved_{card_key}"] = True
                st.session_state.pop("pending_save", None)
                st.session_state.pop(f"show_db_dialog_{card_key}", None)
                st.rerun()
            else:
                st.error(f"Save failed: {_sv.text[:200]}")
        else:
            st.info(f"**'{title}'** is already in the DB (current: v{_latest_ver}). How would you like to save?")
            _dc1, _dc2, _dc3 = st.columns([2, 2, 1])
            with _dc1:
                if st.button(f"➕  Save as new version  (v{_latest_ver + 1})",
                             key=f"db_new_{card_key}", use_container_width=True, type="primary"):
                    _sv = requests.post(f"{BACKEND}/documents/save", json={
                        "title": title, "doc_type": _type_to_save,
                        "doc_format": _fmt_to_save, "content": _content_to_save,
                        "file_ext": _ext_to_save, "save_mode": "new_version",
                    })
                    if _sv.status_code == 200:
                        _svd = _sv.json()
                        st.success(f"✅ Saved as v{_svd.get('version',1)} (ID #{_svd.get('id')})")
                        st.session_state[f"db_saved_{card_key}"] = True
                        st.session_state.pop("pending_save", None)
                        st.session_state.pop(f"show_db_dialog_{card_key}", None)
                        st.rerun()
                    else:
                        st.error(f"Save failed: {_sv.text[:200]}")
            with _dc2:
                if st.button(f"🔄  Overwrite current  (v{_latest_ver})",
                             key=f"db_ow_{card_key}", use_container_width=True):
                    _sv = requests.post(f"{BACKEND}/documents/save", json={
                        "title": title, "doc_type": _type_to_save,
                        "doc_format": _fmt_to_save, "content": _content_to_save,
                        "file_ext": _ext_to_save, "save_mode": "overwrite",
                        "overwrite_id": _latest_id,
                    })
                    if _sv.status_code == 200:
                        st.success(f"✅ Overwritten — v{_latest_ver} updated")
                        st.session_state[f"db_saved_{card_key}"] = True
                        st.session_state.pop("pending_save", None)
                        st.session_state.pop(f"show_db_dialog_{card_key}", None)
                        st.rerun()
                    else:
                        st.error(f"Save failed: {_sv.text[:200]}")
            with _dc3:
                if st.button("✕  Cancel", key=f"db_cancel_{card_key}", use_container_width=True):
                    st.session_state.pop(f"show_db_dialog_{card_key}", None)
                    st.rerun()

    # ── Version history panel (saved docs only) ────────────────────────────────
    if is_saved and st.session_state.get(f"ver_{card_key}", False):
        ver_res = requests.get(f"{BACKEND}/documents/{doc_id}/versions")
        if ver_res.status_code == 200:
            versions = ver_res.json().get("versions", [])
            with st.expander(f"🕓 Version History — {title} ({len(versions)} version(s))", expanded=True):
                for v in versions:
                    v_id      = v["id"]
                    v_num     = v["version"]
                    v_created = to_ist(v.get("created_at", ""))
                    v_fmt     = v.get("doc_format", doc_format)
                    v_ext     = v.get("file_ext") or ("xlsx" if v_fmt == "excel" else "docx")
                    is_latest = (v_id == doc_id)
                    st.markdown(f"**v{v_num}**{'  ✅ latest' if is_latest else ''} &nbsp; 🕒 {v_created} &nbsp; ID #{v_id}")
                    vc1, vc2 = st.columns([1, 1])
                    with vc1:
                        dl = requests.get(f"{BACKEND}/documents/{v_id}/download")
                        if dl.status_code == 200:
                            v_mime = (
                                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                if v_ext == "xlsx"
                                else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                            safe = title.replace(" ", "_")[:40]
                            st.download_button(
                                label=f"⬇️ Download v{v_num}",
                                data=dl.content,
                                file_name=f"{safe}_v{v_num}.{v_ext}",
                                mime=v_mime,
                                use_container_width=True,
                                key=f"dl_v_{v_id}_{card_key}"
                            )
                    with vc2:
                        if st.button(f"🗑 Delete v{v_num}", key=f"del_v_{v_id}_{card_key}", use_container_width=True):
                            requests.delete(f"{BACKEND}/documents/{v_id}")
                            st.rerun()
                    st.markdown("---")

    # ── View (expanded content) ────────────────────────────────────────────────
    if st.session_state.get(f"view_{card_key}", False):
        # Resolve content
        if in_memory_content is not None:
            view_content = in_memory_content
            view_format  = doc_format
        elif is_saved:
            with st.spinner("Loading document..."):
                det = requests.get(f"{BACKEND}/documents/{doc_id}")
            if det.status_code != 200:
                st.error("Could not load document.")
                return
            view_content = det.json().get("content", {})
            view_format  = det.json().get("doc_format", doc_format)
        else:
            st.warning("No content available.")
            return

        with st.expander(f"📄 {title} — Full Content", expanded=True):
            if view_format == "excel":
                for sheet in view_content.get("sheets", []):
                    sname   = sheet.get("sheet_name", "Sheet")
                    headers = sheet.get("headers", [])
                    rows    = sheet.get("rows", [])
                    st.markdown(f"**{sname}**")
                    if headers and rows:
                        max_cols = max(len(headers), max((len(r) for r in rows), default=0))
                        if len(headers) < max_cols:
                            headers = headers + [f"Col {i+1}" for i in range(len(headers), max_cols)]
                        padded = [r + [""] * max(0, len(headers) - len(r)) for r in rows]
                        padded = [r[:len(headers)] for r in padded]
                        st.dataframe(pd.DataFrame(padded, columns=headers), use_container_width=True, hide_index=True)
            else:
                _order = view_content.get("__section_order__")
                if _order:
                    _meta = {"__section_order__", "__show_headings__"}
                    _items = [(k, view_content[k]) for k in _order if k in view_content]
                    _covered = set(_order) | _meta
                    _items += [(k, v) for k, v in view_content.items() if k not in _covered]
                else:
                    _meta = {"__section_order__", "__show_headings__"}
                    _items = [(k, v) for k, v in view_content.items() if k not in _meta]
                _show_hdgs_v = view_content.get("__show_headings__", True)
                for section_name, section_content in _items:
                    if _show_hdgs_v:
                        st.markdown(f"**{section_name}**")
                    if isinstance(section_content, str):
                        st.write(section_content)
                    elif isinstance(section_content, list):
                        for item in section_content:
                            st.write(f"• {item}")
                    elif isinstance(section_content, dict):
                        for k, v in section_content.items():
                            st.write(f"**{k}:** {v}")
                    st.markdown("---")

    st.markdown("")

# ─────────────────────────────────────────────────────────────────────────────
# Document History Page
# ─────────────────────────────────────────────────────────────────────────────

def _render_history_page():
    with st.spinner("Loading history..."):
        res = requests.get(f"{BACKEND}/documents")

    if res.status_code != 200:
        st.error(f"Could not load history: {res.text}")
        return

    docs = res.json().get("documents", [])
    if not docs:
        st.info("No documents saved yet. Generate a document and save it to DB to see it here.")
        return

    # ── Search + Filter bar ───────────────────────────────────────────────────
    _h_c1, _h_c2, _h_c3 = st.columns([3, 1.5, 1.5])
    with _h_c1:
        _search = st.text_input(
            "Search", placeholder="🔎  Search by title or doc type...",
            key="hist_search", label_visibility="collapsed"
        )
    with _h_c2:
        _fmt_opts = ["All formats", "Word", "Excel"]
        _fmt_filter = st.selectbox("Format", _fmt_opts, key="hist_fmt", label_visibility="collapsed")
    with _h_c3:
        _type_opts = ["All types"] + sorted({d.get("doc_type", "General") or "General" for d in docs})
        _type_filter = st.selectbox("Type", _type_opts, key="hist_type", label_visibility="collapsed")

    # ── Apply filters ─────────────────────────────────────────────────────────
    filtered = docs
    if _search.strip():
        _q = _search.strip().lower()
        filtered = [
            d for d in filtered
            if _q in d["title"].lower()
            or _q in (d.get("doc_type") or "").lower()
        ]
    if _fmt_filter != "All formats":
        filtered = [d for d in filtered if d.get("doc_format", "word").lower() == _fmt_filter.lower()]
    if _type_filter != "All types":
        filtered = [d for d in filtered if (d.get("doc_type") or "General") == _type_filter]

    # ── Results count ─────────────────────────────────────────────────────────
    _count_txt = (
        f"{len(filtered)} of {len(docs)} document(s)"
        if len(filtered) != len(docs)
        else f"{len(docs)} document(s)"
    )
    st.markdown(
        f'<p style="color:#7a94b8;font-size:0.88rem;margin-bottom:8px;">{_count_txt}</p>',
        unsafe_allow_html=True
    )

    if not filtered:
        st.info("No documents match your search.")
        return

    st.markdown("---")

    for doc in filtered:
        _render_doc_card({
            "doc_id":     doc["id"],
            "title":      doc["title"],
            "doc_format": doc.get("doc_format", "word"),
            "doc_type":   doc.get("doc_type", "General"),
            "file_ext":   doc.get("file_ext", "docx"),
            "version":    doc.get("version", 1),
            "created_at": doc.get("created_at", ""),
        })


# ─────────────────────────────────────────────────────────────────────────────
# Sidebar + routing
# ─────────────────────────────────────────────────────────────────────────────

with st.sidebar:
    st.markdown("""
    <div class="sidebar-logo">
        <div class="sidebar-logo-text">⚡ DocWizard</div>
        <div class="sidebar-logo-sub">DocForge Hub · CiteRAG Lab · StateCase Assistant</div>
    </div>
    """, unsafe_allow_html=True)
    st.markdown("---")
    page = st.radio("Navigate", [
        "✏️  Generate Document",
        "📚  Document History",
        "─────────────────",
        "💬  Ask Documents",
        "🖊️  Refine Document",
        "⚖️  Compare",
        "🟰  Evaluate",
        "📋  CiteRAG Log",
        "─────────────────",
        "🤖  Assistant",
        "🎫  My Tickets",
        "📝  Assistant Log",
    ], label_visibility="collapsed")
    st.markdown("---")
    # Show CiteRAG index stats when on a CiteRAG page
    _cr_pages = {"💬  Ask Documents", "⚖️  Compare", "🟰  Evaluate", "📋  CiteRAG Log", "🤖  Assistant"}
    if page in _cr_pages:
        try:
            _cr_status = requests.get(f"{CITERAG_BACKEND}/ingest/status", timeout=2)
            if _cr_status.status_code == 200:
                _cr_docs = _cr_status.json().get("documents", [])
                _cr_chunks = sum(d.get("chunk_count", 0) for d in _cr_docs)
                st.markdown(f"""
                <div style="background:rgba(96,165,250,0.06);border:1px solid rgba(96,165,250,0.2);
                            border-radius:8px;padding:10px 14px;font-size:0.78rem;color:#7a94b8;margin-bottom:8px;">
                    🔬 <strong style="color:#60a5fa">CiteRAG Index</strong><br>
                    {len(_cr_docs)} docs &nbsp;·&nbsp; {_cr_chunks:,} chunks
                </div>
                """, unsafe_allow_html=True)
        except Exception:
            st.caption("⚠️ CiteRAG offline")
    else:
        st.markdown("""
    <div style="padding:12px;background:rgba(245,200,66,0.06);border:1px solid rgba(245,200,66,0.2);
                border-radius:8px;font-size:0.78rem;color:#7a94b8;line-height:1.6;">
        💾 Use the <strong style="color:#f5c842">Save to History</strong> button after generation.
    </div>
    """, unsafe_allow_html=True)
    # Always show CiteRAG index status in sidebar
    try:
        _idx_res = requests.get(f"{CITERAG_BACKEND}/ingest/status", timeout=2)
        if _idx_res.status_code == 200:
            _idx_docs   = _idx_res.json().get("documents", [])
            _idx_chunks = sum(d.get("chunk_count", 0) for d in _idx_docs)
            st.markdown(f"""
            <div style="background:rgba(96,165,250,0.06);border:1px solid rgba(96,165,250,0.2);
                        border-radius:8px;padding:10px 14px;font-size:0.76rem;color:#7a94b8;margin-top:8px;">
                🔬 <strong style="color:#60a5fa">CiteRAG Index</strong><br>
                {len(_idx_docs)} docs &nbsp;·&nbsp; {_idx_chunks:,} chunks<br>
                <span style="font-size:0.68rem;color:#4a6080;">Auto-synced with Notion</span>
            </div>
            """, unsafe_allow_html=True)
    except Exception:
        pass

if page == "📚  Document History":
    st.markdown("""
    <div class="page-header">
        <h1>📚 Document History</h1>
        <p>All your generated documents — browse, preview and download.</p>
    </div>
    """, unsafe_allow_html=True)
    _render_notion_dialogs()
    _render_history_page()
    st.stop()

st.markdown("""
<div class="page-header">
    <h1>⚡ DocWizard </h1>
    <p>Conjure professional Word documents and Excel spreadsheets using AI — in seconds.</p>
</div>
""", unsafe_allow_html=True)

_render_notion_dialogs()


# ─────────────────────────────────────────────────────────────────────────────
# Helpers — Word rendering
# ─────────────────────────────────────────────────────────────────────────────

_MD_TABLE_SEPARATOR_RE = re.compile(
    r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+(?:\s*:?-{3,}:?\s*)\|?\s*$"
)


def _split_markdown_table_cells(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def _parse_markdown_table(table_lines: list[str]) -> tuple[list[str], list[list[str]]] | None:
    if len(table_lines) < 2 or not _MD_TABLE_SEPARATOR_RE.match(table_lines[1]):
        return None

    headers = _split_markdown_table_cells(table_lines[0])
    rows = [_split_markdown_table_cells(line) for line in table_lines[2:] if line.strip()]
    row_widths = [len(row) for row in rows]
    width = max([len(headers)] + row_widths)

    headers = headers + [""] * (width - len(headers))
    rows = [row + [""] * (width - len(row)) for row in rows]
    return headers, rows


def _iter_markdown_blocks(content: str):
    lines = content.splitlines()
    i = 0

    while i < len(lines):
        if not lines[i].strip():
            i += 1
            continue

        if (
            i + 1 < len(lines)
            and "|" in lines[i]
            and _MD_TABLE_SEPARATOR_RE.match(lines[i + 1])
        ):
            table_lines = [lines[i], lines[i + 1]]
            i += 2
            while i < len(lines) and lines[i].strip():
                if "|" not in lines[i]:
                    break
                table_lines.append(lines[i])
                i += 1
            yield ("table", table_lines)
            continue

        text_lines = []
        while i < len(lines) and lines[i].strip():
            if (
                i + 1 < len(lines)
                and "|" in lines[i]
                and _MD_TABLE_SEPARATOR_RE.match(lines[i + 1])
            ):
                break
            text_lines.append(lines[i].strip())
            i += 1
        yield ("text", text_lines)


def render_section_content(content):
    if content is None or content == "":
        st.markdown('<p class="doc-paragraph"><em>No content provided.</em></p>', unsafe_allow_html=True)
        return
    if isinstance(content, str):
        content = content.strip()
        if content.startswith("{") or content.startswith("["):
            try:
                parsed = json.loads(content)
                render_section_content(parsed)
                return
            except Exception:
                pass
        for block_type, block_lines in _iter_markdown_blocks(content):
            if block_type == "table":
                parsed_table = _parse_markdown_table(block_lines)
                if parsed_table:
                    headers, rows = parsed_table
                    st.table(pd.DataFrame(rows, columns=headers))
                    continue

            for para in block_lines:
                if para.startswith(("-", "*", "•", "·")):
                    st.markdown(f'<p class="bullet-item">{para.lstrip("-*•· ").strip()}</p>', unsafe_allow_html=True)
                else:
                    st.markdown(f'<p class="doc-paragraph">{para}</p>', unsafe_allow_html=True)
    elif isinstance(content, list):
        for item in content:
            if isinstance(item, dict):
                render_section_content(item)
            else:
                st.markdown(f'<p class="bullet-item">{str(item).strip().lstrip("-*•· ")}</p>', unsafe_allow_html=True)
    elif isinstance(content, dict):
        for key, val in content.items():
            st.markdown(f'<p class="doc-section-heading" style="font-size:1rem;margin-top:14px">{str(key).replace("_"," ").title()}</p>', unsafe_allow_html=True)
            render_section_content(val)
    else:
        st.markdown(f'<p class="doc-paragraph">{str(content)}</p>', unsafe_allow_html=True)


def render_full_document(title, sections):
    st.markdown('<div class="doc-card">', unsafe_allow_html=True)
    st.markdown(f'<p class="doc-title">📄 {title}</p>', unsafe_allow_html=True)
    # Respect stored section order so it survives JSONB round-trip
    order = sections.get("__section_order__")
    show_headings = sections.get("__show_headings__", True)
    meta_keys = {"__section_order__", "__show_headings__"}
    if order:
        items = [(k, sections[k]) for k in order if k in sections]
        covered = set(order) | meta_keys
        items += [(k, v) for k, v in sections.items() if k not in covered]
    else:
        items = [(k, v) for k, v in sections.items() if k not in meta_keys]

    # ── Table of Contents (only when headings are shown) ──────────────────────
    if show_headings and len(items) > 1:
        toc_rows = "".join(
            f'<div style="display:flex;justify-content:space-between;align-items:baseline;'            f'padding:3px 0;border-bottom:1px dotted rgba(99,179,237,0.2);">'            f'<span style="color:#93c5fd;font-size:0.82rem;font-weight:600;">'            f'{i}. {str(k).replace("_"," ").title()}</span>'            f'<span style="color:#4a6080;font-size:0.75rem;white-space:nowrap;margin-left:8px;">·····</span>'            f'</div>'
            for i, (k, _) in enumerate(items, 1)
        )
        st.markdown(
            f'<div style="background:rgba(30,52,100,0.4);border:1px solid rgba(99,179,237,0.2);'            f'border-radius:8px;padding:14px 18px;margin-bottom:18px;">'            f'<p style="color:#60a5fa;font-size:0.72rem;font-weight:700;letter-spacing:2px;'            f'text-transform:uppercase;margin-bottom:10px;">📋 Table of Contents</p>'            f'{toc_rows}</div>',
            unsafe_allow_html=True
        )

    for section_name, content in items:
        if show_headings:
            st.markdown(f'<p class="doc-section-heading">{str(section_name).replace("_"," ").title()}</p>', unsafe_allow_html=True)
        render_section_content(content)
    st.markdown('</div>', unsafe_allow_html=True)


def flatten_to_text(content) -> str:
    if content is None: return ""
    if isinstance(content, str): return content.strip()
    if isinstance(content, list): return "\n".join(f"• {flatten_to_text(i)}" for i in content)
    if isinstance(content, dict):
        # Skip __section_order__ meta key when flattening
        return "\n\n".join(
            f"{str(k).replace('_',' ').title()}:\n{flatten_to_text(v)}"
            for k, v in content.items() if k not in ("__section_order__", "__show_headings__")
        )
    return str(content)


def _add_toc_to_doc(doc, sec_items):
    """
    Insert a native Word TOC field + manual fallback list.
    The Word TOC field auto-populates page numbers when opened in MS Word.
    The manual list is visible in all viewers (Google Docs, PDF, LibreOffice).
    """
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # ── Native Word TOC field ─────────────────────────────────────────────────
    toc_para = doc.add_paragraph()
    toc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = toc_para.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    # ── Note for user ─────────────────────────────────────────────────────────
    note = doc.add_paragraph()
    note_run = note.add_run(
        "(Open in Microsoft Word and press Ctrl+A then F9 to update page numbers)"
    )
    note_run.font.size   = Pt(9)
    note_run.font.italic = True
    note_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    note.paragraph_format.space_after = Pt(8)

    # ── Manual index (works in all viewers) ───────────────────────────────────
    for i, (section_name, _) in enumerate(sec_items, 1):
        label = str(section_name).replace("_", " ").title()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        rl = p.add_run(f"{i}.  {label}")
        rl.font.size = Pt(11)
        rl.font.color.rgb = RGBColor(0x1E, 0x3A, 0x6E)
        rd = p.add_run(f"  {'.' * max(2, 50 - len(label))}")
        rd.font.size = Pt(9)
        rd.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_page_break()


def create_word_document(title, sections_dict) -> BytesIO:
    doc = Document()
    doc.add_heading(title, 0).alignment = 1
    # Respect stored section order, skip meta keys
    _ord          = sections_dict.get("__section_order__")
    _show_hdgs    = sections_dict.get("__show_headings__", True)
    _meta_keys    = {"__section_order__", "__show_headings__"}
    if _ord:
        _sec_items = [(k, sections_dict[k]) for k in _ord if k in sections_dict]
        _covered = set(_ord) | _meta_keys
        _sec_items += [(k, v) for k, v in sections_dict.items() if k not in _covered]
    else:
        _sec_items = [(k, v) for k, v in sections_dict.items() if k not in _meta_keys]

    # Add TOC page for docs with headings and more than 1 section
    if _show_hdgs and len(_sec_items) > 1:
        _add_toc_to_doc(doc, _sec_items)

    for section_name, content in _sec_items:
        if _show_hdgs:
            doc.add_heading(str(section_name).replace("_", " ").title(), level=1)
        for para in [p.strip() for p in flatten_to_text(content).split("\n") if p.strip()]:
            if para.startswith("•"):
                doc.add_paragraph(para.lstrip("• ").strip(), style="List Bullet")
            else:
                p = doc.add_paragraph()
                p.add_run(para).font.size = Pt(11)
        doc.add_paragraph()
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────────────────────────────────────
# Helpers — Excel rendering
# ─────────────────────────────────────────────────────────────────────────────

def render_excel_document(title, sheets):
    st.markdown(f'<div class="doc-card"><p class="doc-title">📊 {title}</p></div>', unsafe_allow_html=True)
    for sheet in sheets:
        sheet_name  = sheet.get("sheet_name", "Sheet")
        description = sheet.get("description", "")
        headers     = sheet.get("headers", [])
        rows        = sheet.get("rows", [])
        header_rows = set(sheet.get("header_rows", []))
        totals_rows = set(sheet.get("totals_rows", []))
        notes       = sheet.get("notes", "")
        st.markdown(f'<p class="doc-section-heading">📋 {sheet_name}</p>', unsafe_allow_html=True)
        if description:
            st.caption(description)
        if headers and rows:
            max_cols = max(len(headers), max((len(r) for r in rows), default=0))
            if len(headers) < max_cols:
                headers = headers + [f"Col {i+1}" for i in range(len(headers), max_cols)]
            padded = [r + [""] * max(0, len(headers) - len(r)) for r in rows]
            padded = [r[:len(headers)] for r in padded]
            df = pd.DataFrame(padded, columns=headers)
            bold_indices = set()
            for x in list(header_rows or []) + list(totals_rows or []):
                try:
                    bold_indices.add(int(x))
                except (ValueError, TypeError):
                    pass
            def style_row(row):
                base = "background-color:#FFFFFF; color:#000000; font-size:13px;"
                if int(row.name) in bold_indices:
                    return [base + "font-weight:bold;" for _ in row]
                return [base + "font-weight:normal;" for _ in row]
            styled = (
                df.style.apply(style_row, axis=1)
                .set_table_styles([{"selector": "thead th", "props": [
                    ("background-color", "#FFFFFF"), ("color", "#000000"),
                    ("font-weight", "bold"), ("font-size", "13px"),
                    ("border-bottom", "2px solid #cccccc")
                ]}])
            )
            st.dataframe(styled, use_container_width=True, hide_index=True)
        if notes:
            st.caption(f"📝 Notes: {notes}")
        st.markdown("---")


def create_excel_file_from_data(excel_data: dict) -> BytesIO:
    res = requests.post(f"{BACKEND}/export/excel", json=excel_data)
    if res.status_code == 200:
        content_type = res.headers.get("content-type", "")
        if "spreadsheet" in content_type or "octet-stream" in content_type:
            return BytesIO(res.content)
        raise RuntimeError(f"Server returned unexpected response: {res.text[:300]}")
    raise RuntimeError(f"Excel export failed [{res.status_code}]: {res.text[:300]}")


# ─────────────────────────────────────────────────────────────────────────────
# DB save helpers
# ─────────────────────────────────────────────────────────────────────────────

def save_to_db(title: str, doc_type: str, doc_format: str,
               content: dict, file_buf: BytesIO = None, file_ext: str = None,
               save_mode: str = "new_version", overwrite_id: int = None):
    payload = {
        "title": title, "doc_type": doc_type, "doc_format": doc_format,
        "content": content, "file_ext": file_ext,
        "save_mode": save_mode, "overwrite_id": overwrite_id,
    }
    if file_buf:
        file_buf.seek(0)
        payload["file_bytes"] = file_buf.read().hex()
    res = requests.post(f"{BACKEND}/documents/save", json=payload)
    if res.status_code == 200:
        data = res.json()
        doc_id = data.get("id")
        ver    = data.get("version", 1)
        mode   = data.get("mode", "new_version")
        if mode == "overwrite":
            st.toast(f"✅ Overwritten (v{ver})", icon="💾")
        else:
            st.toast(f"✅ Saved as v{ver} (ID #{doc_id})", icon="💾")
        return doc_id
    else:
        st.warning(f"Could not save to history: {res.text[:200]}")
        return None


def version_save_dialog(title: str, doc_type: str, doc_format: str,
                        content: dict, file_buf: BytesIO = None, file_ext: str = None):
    """Check if title exists — show Save as new version / Overwrite if it does."""
    check = requests.post(f"{BACKEND}/documents/check-version", json={"title": title})
    if check.status_code != 200:
        st.warning("Could not check version history.")
        return
    data   = check.json()
    exists = data.get("exists", False)
    if not exists:
        save_to_db(title, doc_type, doc_format, content, file_buf, file_ext, save_mode="new_version")
        return
    latest_ver = data.get("latest_version", 1)
    latest_id  = data.get("latest_id")
    st.info(f"**'{title}'** is already saved (current: v{latest_ver}). Choose how to save:")
    col1, col2 = st.columns(2)
    with col1:
        if st.button(f"📄 Save as new version (v{latest_ver + 1})", use_container_width=True):
            save_to_db(title, doc_type, doc_format, content, file_buf, file_ext, save_mode="new_version")
            st.session_state.pop("pending_save", None)
            st.rerun()
    with col2:
        if st.button(f"🔄 Overwrite current (v{latest_ver})", use_container_width=True):
            save_to_db(title, doc_type, doc_format, content, file_buf, file_ext,
                       save_mode="overwrite", overwrite_id=latest_id)
            st.session_state.pop("pending_save", None)
            st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
if page == "✏️  Generate Document":
    # STEP 1: Plan
    # ─────────────────────────────────────────────────────────────────────────────

    st.markdown('<div class="step-label">Step 1 — Describe Your Document</div>', unsafe_allow_html=True)
    user_prompt = st.text_input(
        "Describe the document you want to generate",
        placeholder="e.g. Balance Sheet for FY2024  •  SOP for lead qualification  •  Product Proposal",
        label_visibility="collapsed"
    )

    if st.button("⚡ Plan Document"):
        if not user_prompt.strip():
            st.warning("Please enter a document description.")
        else:
            with st.spinner("Planning document structure..."):
                res = requests.post(f"{BACKEND}/plan", json={"prompt": user_prompt})
            if res.status_code == 200:
                plan = res.json()
                st.session_state["title"]         = plan["title"]
                st.session_state["sections"]      = plan["sections"]
                st.session_state["doc_format"]    = plan.get("doc_format", "word")
                st.session_state["doc_type"]      = plan.get("doc_type", "General")
                st.session_state["show_headings"] = plan.get("show_headings", True)
                for k in ["questions", "generated_sections", "excel_data"]:
                    st.session_state.pop(k, None)
                st.success("Document structure generated!")
            else:
                st.error(f"Error: {res.text}")


    # ─────────────────────────────────────────────────────────────────────────────
    # STEP 2: Show Plan
    # ─────────────────────────────────────────────────────────────────────────────

    if "sections" in st.session_state:
        doc_format = st.session_state.get("doc_format", "word")
        badge = "excel" if doc_format == "excel" else "word"
        label = "📊 Excel / Spreadsheet" if doc_format == "excel" else "📄 Word Document"

        st.markdown('<div class="step-label">Step 2 — Review Structure</div>', unsafe_allow_html=True)
        col1, col2, col3 = st.columns([3, 1, 1])
        with col1:
            if st.session_state.get("editing_title"):
                new_title_val = st.text_input(
                    "Edit title", value=st.session_state["title"],
                    key="title_edit_input", label_visibility="collapsed"
                )
            else:
                st.markdown(f"**Title:** {st.session_state['title']}")
        with col2:
            st.markdown(f'<span class="format-badge-{badge}">{label}</span>', unsafe_allow_html=True)
        with col3:
            if st.session_state.get("editing_title"):
                if st.button("✅", key="title_save_btn", help="Save title"):
                    if new_title_val.strip():
                        st.session_state["title"] = new_title_val.strip()
                    st.session_state["editing_title"] = False
                    st.rerun()
            else:
                if st.button("✏️", key="title_edit_btn", help="Edit title"):
                    st.session_state["editing_title"] = True
                    st.rerun()

        for i, sec in enumerate(st.session_state["sections"]):
            col_sec, col_del = st.columns([10, 1])
            with col_sec:
                st.write(f"- {sec}")
            with col_del:
                if st.button("✕", key=f"del_sec_{i}", help=f"Remove '{sec}'"):
                    st.session_state["sections"].pop(i)
                    for k in ["questions", "generated_sections", "excel_data"]:
                        st.session_state.pop(k, None)
                    st.rerun()

        with st.expander("➕ Add a section", expanded=False):
            new_sec = st.text_input("Section name", placeholder="e.g. Risk Analysis, Appendix...", key="new_section_input")
            insert_pos = st.selectbox("Insert position", options=["At the end"] + [f"Before: {s}" for s in st.session_state["sections"]], key="new_section_pos")
            if st.button("Add Section", key="add_section_btn"):
                if new_sec.strip():
                    sec_name = new_sec.strip()
                    if insert_pos == "At the end":
                        st.session_state["sections"].append(sec_name)
                    else:
                        ref = insert_pos.replace("Before: ", "")
                        idx = st.session_state["sections"].index(ref)
                        st.session_state["sections"].insert(idx, sec_name)
                    for k in ["questions", "generated_sections", "excel_data"]:
                        st.session_state.pop(k, None)
                    st.success(f"Added '{sec_name}'. Regenerate questions to include it.")
                    st.rerun()
                else:
                    st.warning("Please enter a section name.")

        if st.button("📋 Generate Questions"):
            with st.spinner("Generating questions..."):
                res = requests.post(f"{BACKEND}/questions", json={"title": st.session_state["title"], "sections": st.session_state["sections"]})
            if res.status_code == 200:
                st.session_state["questions"] = res.json()
                st.session_state["q_page"]    = 0   # start at page 1
                st.session_state["answers"]   = {}  # clear any previous answers
                for k in ["generated_sections", "excel_data"]:
                    st.session_state.pop(k, None)
                st.success("Questions generated!")
            else:
                st.error(f"Error: {res.text}")


    # ─────────────────────────────────────────────────────────────────────────────
    # STEP 3: Collect Answers — paginated, one section per page
    # ─────────────────────────────────────────────────────────────────────────────

    if "questions" in st.session_state:
        st.markdown('<div class="step-label">Step 3 — Provide Details (optional)</div>', unsafe_allow_html=True)

        # Build ordered list of sections that have questions
        _q_sections = [s for s in st.session_state.get("sections", [])
                       if s in st.session_state["questions"]
                       and isinstance(st.session_state["questions"][s], list)]
        # Append any sections in questions dict not in the plan order
        for _s in st.session_state["questions"]:
            if _s not in _q_sections and isinstance(st.session_state["questions"][_s], list):
                _q_sections.append(_s)

        _total_pages = len(_q_sections)

        # Initialise page index in session state
        if "q_page" not in st.session_state or st.session_state.get("q_page", 0) >= _total_pages:
            st.session_state["q_page"] = 0
        _cur_page = st.session_state["q_page"]

        # ── Progress bar + section header ────────────────────────────────────────
        if _total_pages > 0:
            _progress = (_cur_page + 1) / _total_pages
            st.progress(_progress)
            st.markdown(
                f'<div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:6px;">' +
                f'<span style="font-family:IBM Plex Mono,monospace;font-size:0.72rem;color:#5a7a9e;letter-spacing:0.08em;">' +
                f'SECTION {_cur_page + 1} OF {_total_pages}</span>' +
                f'<span style="font-family:IBM Plex Mono,monospace;font-size:0.72rem;color:#f5c842;">{_q_sections[_cur_page]}</span>' +
                f'</div>',
                unsafe_allow_html=True
            )

            # ── Section title card ────────────────────────────────────────────────
            _sec_name  = _q_sections[_cur_page]
            _sec_qs    = st.session_state["questions"][_sec_name]

            st.markdown(
                f'<div style="background:linear-gradient(135deg,#0d1836,#111f38);' +
                f'border:1px solid #1e3055;border-left:4px solid #f5c842;' +
                f'border-radius:10px;padding:16px 22px;margin-bottom:20px;">' +
                f'<p style="font-size:1.1rem;font-weight:700;color:#ffffff;margin:0;">📋 {_sec_name}</p>' +
                f'<p style="font-size:0.78rem;color:#7a94b8;margin:4px 0 0;">' +
                f'{len(_sec_qs)} question(s) — answers are optional but improve generation quality</p>' +
                f'</div>',
                unsafe_allow_html=True
            )

            # ── Questions for this section ────────────────────────────────────────
            # Preserve answers across page turns using session state
            if "answers" not in st.session_state:
                st.session_state["answers"] = {}

            for _qi, _q in enumerate(_sec_qs):
                _key = f"{_sec_name}__{_q}"
                _cur_val = st.session_state["answers"].get(_q, "")
                _new_val = st.text_area(
                    f"**Q{_qi + 1}.** {_q}",
                    value=_cur_val,
                    key=f"qa_{_key}",
                    placeholder="Leave blank to let AI decide…",
                    height=90,
                )
                st.session_state["answers"][_q] = _new_val

            # ── Navigation buttons ────────────────────────────────────────────────
            st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
            _nav_cols = st.columns([1, 2, 1])

            with _nav_cols[0]:
                if _cur_page > 0:
                    if st.button("← Previous", use_container_width=True, key="q_prev"):
                        st.session_state["q_page"] -= 1
                        st.rerun()

            with _nav_cols[1]:
                # Dot indicators
                _dots = ""
                for _di in range(_total_pages):
                    _color = "#f5c842" if _di == _cur_page else "#1e3055"
                    _size  = "10px"    if _di == _cur_page else "8px"
                    _dots += (f'<span style="display:inline-block;width:{_size};height:{_size};' +
                              f'border-radius:50%;background:{_color};margin:0 4px;' +
                              f'vertical-align:middle;"></span>')
                st.markdown(
                    f'<div style="text-align:center;padding:8px 0;">{_dots}</div>',
                    unsafe_allow_html=True
                )

            with _nav_cols[2]:
                if _cur_page < _total_pages - 1:
                    if st.button("Next →", use_container_width=True, key="q_next", type="primary"):
                        st.session_state["q_page"] += 1
                        st.rerun()

        # ── Add another section (always visible at bottom) ────────────────────────
        with st.expander("➕ Add another section", expanded=False):
            new_sec3 = st.text_input("New section name", placeholder="e.g. Competitive Analysis...", key="new_section_input_s3")
            insert_pos3 = st.selectbox("Insert position", options=["At the end"] + [f"Before: {s}" for s in st.session_state.get("sections", [])], key="new_section_pos_s3")
            if st.button("Add & Regenerate Questions", key="add_section_btn_s3"):
                if new_sec3.strip():
                    sec_name3 = new_sec3.strip()
                    sections  = st.session_state.get("sections", [])
                    if insert_pos3 == "At the end":
                        sections.append(sec_name3)
                    else:
                        ref3 = insert_pos3.replace("Before: ", "")
                        idx3 = sections.index(ref3)
                        sections.insert(idx3, sec_name3)
                    st.session_state["sections"] = sections
                    with st.spinner(f"Generating questions for '{sec_name3}'..."):
                        res = requests.post(f"{BACKEND}/questions", json={"title": st.session_state["title"], "sections": sections})
                    if res.status_code == 200:
                        st.session_state["questions"] = res.json()
                        st.session_state["q_page"] = 0   # reset to first page
                        st.success(f"Added '{sec_name3}' and regenerated questions.")
                        st.rerun()
                    else:
                        st.error(f"Could not regenerate questions: {res.text}")
                else:
                    st.warning("Please enter a section name.")

        # ── Generate button — only on last page (or if no sections) ──────────────
        _answers = st.session_state.get("answers", {})
        _on_last  = (_total_pages == 0 or _cur_page == _total_pages - 1)
        if _on_last:
            st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)
            st.divider()

        if _on_last and st.button("🚀 Generate Document", use_container_width=True, type="primary"):
            answers = _answers
            with st.spinner("Starting generation job..."):
                res = requests.post(f"{BACKEND}/generate", json={
                    "title": st.session_state["title"], "sections": st.session_state["sections"],
                    "answers": answers, "doc_format": st.session_state.get("doc_format", "word"),
                    "show_headings": st.session_state.get("show_headings", True),
                })
            if res.status_code != 200:
                st.error(f"Error: {res.text}")
            else:
                raw        = res.json()
                doc_format = st.session_state.get("doc_format", "word")
                doc_type   = st.session_state.get("doc_type", "General")

                if doc_format == "excel":
                    if isinstance(raw, dict) and "sheets" not in raw:
                        for v in raw.values():
                            if isinstance(v, dict) and "sheets" in v:
                                raw = v
                                break
                    if isinstance(raw, dict) and "title" not in raw:
                        raw["title"] = st.session_state.get("title", "Document")
                    raw.pop("doc_type", None)
                    st.session_state["excel_data"] = raw
                    st.session_state.pop("generated_sections", None)
                    st.session_state["pending_save"] = {
                        "title": st.session_state["title"], "doc_type": doc_type,
                        "doc_format": "excel", "content": raw, "file_ext": "xlsx",
                    }
                else:
                    if isinstance(raw, dict) and list(raw.keys()) == ["Document"]:
                        raw = raw["Document"]
                    if isinstance(raw, dict) and list(raw.keys()) == ["sections"]:
                        raw = raw["sections"]
                    sections_dict = raw if isinstance(raw, dict) else {"Content": str(raw)}
                    # Preserve section order + show_headings flag so both survive JSONB round-trip
                    planned_order = st.session_state.get("sections", [])
                    show_hdgs     = sections_dict.pop("__show_headings__",
                                        st.session_state.get("show_headings", True))
                    sections_dict["__show_headings__"] = show_hdgs
                    if planned_order:
                        sections_dict["__section_order__"] = [
                            s for s in planned_order if s in sections_dict
                        ] + [s for s in sections_dict if s not in planned_order
                             and s not in ("__section_order__", "__show_headings__")]
                    st.session_state["generated_sections"] = sections_dict
                    st.session_state.pop("excel_data", None)
                    st.session_state["pending_save"] = {
                        "title": st.session_state["title"], "doc_type": doc_type,
                        "doc_format": "word", "content": sections_dict, "file_ext": "docx",
                    }
                st.success("Document generated!")
                st.rerun()


    # ─────────────────────────────────────────────────────────────────────────────
    # STEP 4A: Display Word Document
    # ─────────────────────────────────────────────────────────────────────────────

    if "generated_sections" in st.session_state:
        st.divider()
        st.markdown('<div class="step-label">Generated Document</div>', unsafe_allow_html=True)
        render_full_document(st.session_state["title"], st.session_state["generated_sections"])

        st.divider()
        st.markdown('<div class="step-label">Review & Refine</div>', unsafe_allow_html=True)
        for section, content in [(k, v) for k, v in st.session_state["generated_sections"].items() if k not in ("__section_order__", "__show_headings__")]:
            with st.expander(f"✏️ Edit: {section}", expanded=False):
                feedback = st.text_area(f"Suggest changes for '{section}'", key=f"feedback_{section}")
                if st.button(f"Update '{section}'", key=f"btn_{section}"):
                    with st.spinner(f"Refining {section}..."):
                        res = requests.post(f"{BACKEND}/refine-section", json={
                            "section_name": section, "original_text": flatten_to_text(content),
                            "feedback": feedback, "doc_format": "word"
                        })
                    if res.status_code == 200:
                        result  = res.json()
                        updated = result.get("updated_text", "")
                        if isinstance(updated, (dict, list)):
                            updated = flatten_to_text(updated)
                        st.session_state["generated_sections"] = {**st.session_state["generated_sections"], section: str(updated)}
                        st.success(f"'{section}' updated!")
                        st.rerun()
                    elif res.status_code == 429:
                        st.warning("⏳ Too many refinements — please wait a moment before trying again.")
                    else:
                        st.error(f"Error: {res.text}")

        st.divider()
        st.markdown('<div class="step-label">Actions</div>', unsafe_allow_html=True)

        _word_buf = create_word_document(st.session_state["title"], st.session_state["generated_sections"])
        _render_doc_card(
            doc={
                "doc_id":     None,
                "title":      st.session_state["title"],
                "doc_format": "word",
                "doc_type":   st.session_state.get("doc_type", "General"),
                "file_ext":   "docx",
                "version":    1,
                "created_at": "",
            },
            in_memory_content = st.session_state["generated_sections"],
            in_memory_buf     = _word_buf,
        )

        st.markdown("---")
        if st.button("✨ Create Another Document", use_container_width=True, key="new_doc_word"):
            for key in ["title", "sections", "doc_format", "doc_type",
                        "questions", "answers", "q_page", "generated_sections",
                        "pending_save", "saved_doc_id", "saved_doc_ver",
                        "notion_page_id_word", "notion_page_id_excel",
                        "show_version_dialog_word"]:
                st.session_state.pop(key, None)
            st.rerun()


    # ─────────────────────────────────────────────────────────────────────────────
    # STEP 4B: Display Excel Document
    # ─────────────────────────────────────────────────────────────────────────────

    if "excel_data" in st.session_state:
        excel_data = st.session_state["excel_data"]
        sheets     = excel_data.get("sheets", [])

        st.divider()
        st.markdown('<div class="step-label">Generated Spreadsheet</div>', unsafe_allow_html=True)
        render_excel_document(st.session_state["title"], sheets)

        st.divider()
        st.markdown('<div class="step-label">Review & Refine Sheets</div>', unsafe_allow_html=True)
        for i, sheet in enumerate(sheets):
            sheet_name = sheet.get("sheet_name", f"Sheet {i+1}")
            with st.expander(f"✏️ Edit: {sheet_name}", expanded=False):
                headers = sheet.get("headers", [])
                rows    = sheet.get("rows", [])
                if headers and rows:
                    max_cols = max(len(headers), max((len(r) for r in rows), default=0))
                    if len(headers) < max_cols:
                        headers = headers + [f"Col {j+1}" for j in range(len(headers), max_cols)]
                    padded = [r + [""] * max(0, len(headers) - len(r)) for r in rows]
                    padded = [r[:len(headers)] for r in padded]
                    st.dataframe(pd.DataFrame(padded, columns=headers), use_container_width=True, hide_index=True)
                feedback = st.text_area(
                    f"Suggest changes for '{sheet_name}'",
                    placeholder="Be specific. e.g. Update Cash value to 250000 • Add row: Deferred Tax = 75000",
                    key=f"excel_feedback_{i}", height=100
                )
                if st.button(f"Update '{sheet_name}'", key=f"excel_btn_{i}"):
                    with st.spinner(f"Refining {sheet_name}..."):
                        res = requests.post(f"{BACKEND}/refine-section", json={
                            "section_name": sheet_name, "current_data": sheet,
                            "feedback": feedback, "doc_format": "excel"
                        })
                    if res.status_code == 200:
                        updated_sheet = res.json().get("updated_sheet", sheet)
                        if not updated_sheet.get("rows"):
                            st.error("Refinement returned empty rows — keeping original. Try rephrasing your feedback.")
                        else:
                            updated_sheets = list(sheets)
                            updated_sheets[i] = updated_sheet
                            st.session_state["excel_data"] = {**st.session_state["excel_data"], "sheets": updated_sheets}
                            st.success(f"'{sheet_name}' updated!")
                            st.rerun()
                    else:
                        st.error(f"Refinement failed [{res.status_code}]: {res.text[:300]}")

        st.divider()
        st.markdown('<div class="step-label">Actions</div>', unsafe_allow_html=True)

        try:
            _xlsx_buf = create_excel_file_from_data(st.session_state["excel_data"])
        except Exception as e:
            st.error(f"Excel export error: {e}")
            st.info("Make sure `openpyxl` is installed: `pip install openpyxl`")
            _xlsx_buf = None

        _render_doc_card(
            doc={
                "doc_id":     None,
                "title":      st.session_state["title"],
                "doc_format": "excel",
                "doc_type":   st.session_state.get("doc_type", "General"),
                "file_ext":   "xlsx",
                "version":    1,
                "created_at": "",
            },
            in_memory_content = st.session_state["excel_data"],
            in_memory_buf     = _xlsx_buf,
        )

        st.markdown("---")
        if st.button("✨ Create Another Document", use_container_width=True, key="new_doc_excel"):
            for key in ["title", "sections", "doc_format", "doc_type",
                        "questions", "answers", "q_page", "generated_sections", "excel_data",
                        "pending_save", "saved_doc_id", "saved_doc_ver",
                        "notion_page_id_word", "notion_page_id_excel",
                        "show_version_dialog_excel"]:
                st.session_state.pop(key, None)
            st.rerun()
    st.stop()



# ─────────────────────────────────────────────────────────────────────────────
# CITERAG PAGES — shared helpers
# ─────────────────────────────────────────────────────────────────────────────

def _cr_score_color(score: float) -> str:
    if score >= 0.8:  return "#22c55e"
    if score >= 0.6:  return "#f59e0b"
    return "#ef4444"


def _cr_render_sources(sources: list, expanded: bool = True):
    if not sources:
        st.caption("No sources retrieved.")
        return
    with st.expander(f"📚 {len(sources)} source(s) used", expanded=expanded):
        for i, src in enumerate(sources, 1):
            score      = src.get("score", 0)
            color      = _cr_score_color(score)
            notion_url = src.get("notion_url", "")
            link       = (f' &nbsp;<a href="{notion_url}" target="_blank" '
                          f'style="color:#60a5fa;font-size:0.7rem;">Open in Notion ↗</a>'
                          if notion_url else "")
            chunk = src.get("chunk_text", "")
            preview = chunk[:400] + ("..." if len(chunk) > 400 else "")
            st.markdown(f"""
            <div style="background:#111827;border:1px solid #1e3a5f;border-left:3px solid #1a56db;
                        border-radius:6px;padding:12px 16px;margin-bottom:8px;font-size:0.84rem;">
                <div style="font-weight:600;color:#60a5fa;font-size:0.82rem;">
                    [Source {i}] {src.get('doc_title','')}
                    <span style="color:#64748b;font-weight:400;"> v{src.get('version',1)}</span>
                    {link}
                </div>
                <div style="color:#64748b;font-size:0.72rem;margin:3px 0 8px;">
                    📂 {src.get('section_name','')} &nbsp;|&nbsp;
                    <span style="background:{color}22;color:{color};border:1px solid {color}55;
                                 padding:2px 7px;border-radius:4px;font-weight:700;font-size:0.7rem;">
                        {score:.0%} match
                    </span>
                </div>
                <div style="color:#94a3b8;line-height:1.5;">{preview}</div>
            </div>
            """, unsafe_allow_html=True)


def _cr_fetch_docs():
    try:
        res = requests.get(f"{CITERAG_BACKEND}/documents", timeout=5)
        if res.status_code == 200:
            return res.json()
    except Exception:
        pass
    return {"titles": [], "industries": []}


def _cr_check_backend():
    try:
        res = requests.get(f"{CITERAG_BACKEND}/health", timeout=3)
        return res.status_code == 200
    except Exception:
        return False


# ─────────────────────────────────────────────────────────────────────────────
# PAGE: Assistant
# ─────────────────────────────────────────────────────────────────────────────

if page == "🤖  Assistant":
    st.markdown("""
    <div class="page-header">
        <h1>🤖 DocWizard Assistant</h1>
        <p>Stateful AI assistant — asks clarifying questions, answers from your docs, and raises tickets when it can't.</p>
    </div>
    """, unsafe_allow_html=True)

    # ── Index status check ────────────────────────────────────────────────────
    try:
        _idx_check = requests.get(f"{CITERAG_BACKEND}/ingest/status", timeout=2)
        if _idx_check.status_code == 200:
            _idx_docs = _idx_check.json().get("documents", [])
            if not _idx_docs:
                st.warning(
                    "⚠️ **No documents indexed yet.** The assistant retrieves answers from "
                    "your Notion documents. Push documents to Notion first, then they will "
                    "be auto-ingested. Until then, all questions will create tickets."
                )
            else:
                _idx_chunks = sum(d.get("chunk_count", 0) for d in _idx_docs)
                st.caption(f"📚 {len(_idx_docs)} document(s) · {_idx_chunks:,} chunks indexed and available")
    except Exception:
        pass

    # ── Session controls ──────────────────────────────────────────────────────
    _a_col1, _a_col2, _a_col3 = st.columns([3, 1.5, 1.5])
    with _a_col1:
        _industry_opts = ["None", "Finance", "HR", "Legal", "Tech", "Operations", "Sales", "General"]
        _a_industry = st.selectbox("Industry context", _industry_opts, key="asst_industry",
                                   label_visibility="collapsed")
    with _a_col2:
        if st.button("🔄 New Conversation", use_container_width=True, key="asst_new"):
            for k in ["asst_thread_id", "asst_messages"]:
                st.session_state.pop(k, None)
            st.rerun()
    with _a_col3:
        st.markdown(
            f'<div style="font-size:0.75rem;color:#4a6080;padding-top:8px;">'            f'Thread: <code style="color:#60a5fa;">{st.session_state.get("asst_thread_id","new")[:8]}...</code></div>',
            unsafe_allow_html=True
        )

    # ── Chat history display ──────────────────────────────────────────────────
    _asst_msgs = st.session_state.get("asst_messages", [])
    for _msg in _asst_msgs:
        _role = _msg.get("role", "user")
        _content = _msg.get("content", "")
        _sources = _msg.get("sources", [])
        _tkt_url = _msg.get("ticket_url", "")

        if _role == "user":
            st.markdown(
                f'<div style="background:rgba(30,58,138,0.3);border:1px solid rgba(99,179,237,0.2);'                f'border-radius:8px;padding:12px 16px;margin:8px 0;text-align:right;">'                f'<span style="color:#93c5fd;font-size:0.85rem;">🧑 You</span><br>'                f'<span style="color:#e2e8f0;">{_content}</span></div>',
                unsafe_allow_html=True
            )
        else:
            st.markdown(
                f'<div style="background:rgba(15,28,50,0.8);border:1px solid rgba(99,179,237,0.15);'                f'border-radius:8px;padding:14px 18px;margin:8px 0;">'                f'<span style="color:#60a5fa;font-size:0.85rem;">🤖 Assistant</span><br>'                f'<span style="color:#e2e8f0;line-height:1.7;">{_content.replace(chr(10), "<br>")}</span></div>',
                unsafe_allow_html=True
            )
            if _sources:
                with st.expander(f"📚 {len(_sources)} source(s) used", expanded=False):
                    _cr_render_sources(_sources)
            if _tkt_url:
                st.markdown(
                    f'<a href="{_tkt_url}" target="_blank" style="font-size:0.8rem;color:#f5c842;">'                    f'🎫 View ticket in Notion ↗</a>',
                    unsafe_allow_html=True
                )

    # ── Input box ─────────────────────────────────────────────────────────────
    st.markdown("")
    with st.form("asst_input_form", clear_on_submit=True):
        _a_input_col, _a_send_col = st.columns([8, 1])
        with _a_input_col:
            _user_input = st.text_input(
                "Message",
                placeholder="Ask anything about your documents...",
                label_visibility="collapsed"
            )
        with _a_send_col:
            _send = st.form_submit_button("⚡", use_container_width=True,
                                          help="Send message")

    if _send and _user_input.strip():
        with st.spinner("Thinking..."):
            try:
                _payload = {
                    "message":   _user_input.strip(),
                    "thread_id": st.session_state.get("asst_thread_id"),
                    "filters":   {},
                }
                if _a_industry and _a_industry != "None":
                    _payload["industry"] = _a_industry

                _res = requests.post(f"{BACKEND}/assistant/chat", json=_payload, timeout=60)

                if _res.status_code == 200:
                    _data = _res.json()
                    st.session_state["asst_thread_id"] = _data["thread_id"]
                    st.session_state["asst_messages"]  = _data["messages"]
                    st.rerun()
                else:
                    st.error(f"Assistant error: {_res.text[:200]}")
            except Exception as _ae:
                st.error(f"Connection error: {_ae}")

    # ── Suggested prompts ─────────────────────────────────────────────────────
    if not _asst_msgs:
        st.markdown("")
        st.markdown('<p style="color:#4a6080;font-size:0.8rem;">Try asking:</p>', unsafe_allow_html=True)
        _suggestions = [
            "What is the incident response procedure in our policies?",
            "Summarise the key clauses in our vendor agreements.",
            "What are employee leave entitlements?",
            "Compare our SOW and MSA documents.",
        ]
        _s_cols = st.columns(len(_suggestions))
        for _sc, _sq in zip(_s_cols, _suggestions):
            with _sc:
                if st.button(_sq[:35] + "...", key=f"asst_sugg_{_sq[:10]}", use_container_width=True):
                    st.session_state["asst_input_prefill"] = _sq
                    st.rerun()

        # Handle prefill
        if "asst_input_prefill" in st.session_state:
            _pf = st.session_state.pop("asst_input_prefill")
            with st.spinner("Thinking..."):
                _pf_res = requests.post(f"{BACKEND}/assistant/chat", json={
                    "message":   _pf,
                    "thread_id": st.session_state.get("asst_thread_id"),
                }, timeout=60)
            if _pf_res.status_code == 200:
                _pf_data = _pf_res.json()
                st.session_state["asst_thread_id"] = _pf_data["thread_id"]
                st.session_state["asst_messages"]  = _pf_data["messages"]
                st.rerun()

    st.stop()


# ─────────────────────────────────────────────────────────────────────────────
# PAGE: My Tickets
# ─────────────────────────────────────────────────────────────────────────────

if page == "🎫  My Tickets":
    st.markdown("""
    <div class="page-header">
        <h1>🎫 My Tickets</h1>
        <p>Support tickets raised when the assistant couldn't find an answer — track status and resolutions.</p>
    </div>
    """, unsafe_allow_html=True)

    _tk_col1, _tk_col2 = st.columns([2, 1])
    with _tk_col1:
        _tk_limit = st.slider("Tickets to load", 5, 50, 20, key="tk_limit")
    with _tk_col2:
        if st.button("🔄 Refresh", key="tk_refresh"):
            st.rerun()

    with st.spinner("Loading tickets..."):
        _tk_res = requests.get(f"{BACKEND}/assistant/tickets", params={"limit": _tk_limit})

    if _tk_res.status_code != 200:
        st.error(f"Could not load tickets: {_tk_res.text[:200]}")
        st.stop()

    _tickets = _tk_res.json().get("tickets", [])

    if not _tickets:
        st.info("No tickets yet. Tickets are created automatically when the assistant can't find an answer.")
        st.stop()

    # ── Status filter ─────────────────────────────────────────────────────────
    _tk_status_filter = st.selectbox(
        "Filter by status", ["All", "Open", "In Progress", "Resolved"],
        key="tk_status_filter", label_visibility="collapsed"
    )
    if _tk_status_filter != "All":
        _tickets = [t for t in _tickets if t.get("status") == _tk_status_filter]

    st.markdown(f'<p style="color:#7a94b8;font-size:0.88rem;">{len(_tickets)} ticket(s)</p>', unsafe_allow_html=True)
    st.markdown("---")

    _status_colors = {"Open": "#ef4444", "In Progress": "#f59e0b", "Resolved": "#22c55e"}
    _priority_colors = {"High": "#ef4444", "Medium": "#f59e0b", "Low": "#6b7280"}

    for _tk in _tickets:
        _sc = _status_colors.get(_tk["status"], "#6b7280")
        _pc = _priority_colors.get(_tk["priority"], "#6b7280")
        with st.expander(
            f"🎫 {_tk['title'][:80]}  —  {_tk['created_at']}",
            expanded=False
        ):
            _badge_row = (
                f'<span style="padding:2px 10px;border-radius:12px;font-size:0.7rem;font-weight:700;'                f'background:{_sc}22;color:{_sc};border:1px solid {_sc}55;margin-right:6px;">{_tk["status"]}</span>'                f'<span style="padding:2px 10px;border-radius:12px;font-size:0.7rem;font-weight:700;'                f'background:{_pc}22;color:{_pc};border:1px solid {_pc}55;margin-right:6px;">{_tk["priority"]}</span>'                f'<span style="padding:2px 10px;border-radius:12px;font-size:0.7rem;font-weight:700;'                f'background:rgba(99,179,237,0.1);color:#60a5fa;border:1px solid rgba(99,179,237,0.3);">{_tk["type"]}</span>'
            )
            st.markdown(_badge_row, unsafe_allow_html=True)
            st.markdown("")

            if _tk.get("question"):
                st.markdown(f"**Question:** {_tk['question']}")
            if _tk.get("summary"):
                st.markdown(f"**Summary:** {_tk['summary']}")
            if _tk.get("assigned_to") and _tk["assigned_to"] != "Unassigned":
                st.markdown(f"**Assigned to:** {_tk['assigned_to']}")

            _tk_c1, _tk_c2, _tk_c3 = st.columns([2, 2, 2])
            with _tk_c1:
                _new_status = st.selectbox(
                    "Update status", ["Open", "In Progress", "Resolved"],
                    index=["Open", "In Progress", "Resolved"].index(_tk["status"]) if _tk["status"] in ["Open", "In Progress", "Resolved"] else 0,
                    key=f"tk_status_{_tk['ticket_id']}"
                )
            with _tk_c2:
                if st.button("💾 Update", key=f"tk_update_{_tk['ticket_id']}", use_container_width=True):
                    _upd = requests.patch(
                        f"{BACKEND}/assistant/tickets/{_tk['ticket_id']}/status",
                        json={"status": _new_status}
                    )
                    if _upd.status_code == 200:
                        st.success(f"Updated to '{_new_status}'")
                        st.rerun()
                    else:
                        st.error("Update failed")
            with _tk_c3:
                if _tk.get("url"):
                    _tk_url = _tk["url"]
                    st.markdown(
                        f'<a href="{_tk_url}" target="_blank" '
                        f'style="font-size:0.82rem;color:#60a5fa;">Open in Notion ↗</a>',
                        unsafe_allow_html=True
                    )
    st.stop()


# ─────────────────────────────────────────────────────────────────────────────
# PAGE: Assistant Log
# ─────────────────────────────────────────────────────────────────────────────

if page == "📝  Assistant Log":
    st.markdown("""
    <div class="page-header">
        <h1>📝 Assistant Log</h1>
        <p>Every assistant conversation — questions, answers, sources and outcomes — saved to Notion.</p>
    </div>
    """, unsafe_allow_html=True)

    _al_c1, _al_c2 = st.columns([2, 1])
    with _al_c1:
        _al_limit = st.slider("Entries to load", 10, 100, 30, key="al_limit")
    with _al_c2:
        if st.button("🔄 Refresh", key="al_refresh"):
            st.rerun()

    with st.spinner("Loading assistant log..."):
        _al_res = requests.get(f"{BACKEND}/assistant/log", params={"limit": _al_limit}, timeout=10)

    if _al_res.status_code != 200:
        st.error(f"Could not load assistant log: {_al_res.text[:200]}")
        st.stop()

    _al_entries = _al_res.json().get("entries", [])

    if not _al_entries:
        st.info("No entries yet. Start a conversation in the 🤖 Assistant tab.")
        st.stop()

    # Filter bar
    _al_f1, _al_f2 = st.columns(2)
    with _al_f1:
        _al_outcome_filter = st.selectbox(
            "Filter by outcome", ["All", "Answered", "Ticket Created", "Clarification Asked"],
            key="al_outcome", label_visibility="collapsed"
        )
    with _al_f2:
        _al_intent_filter = st.selectbox(
            "Filter by intent", ["All", "question", "generate", "compare", "clarification", "chitchat"],
            key="al_intent", label_visibility="collapsed"
        )

    if _al_outcome_filter != "All":
        _al_entries = [e for e in _al_entries if e.get("outcome") == _al_outcome_filter]
    if _al_intent_filter != "All":
        _al_entries = [e for e in _al_entries if e.get("intent") == _al_intent_filter]

    st.markdown(f'<p style="color:#7a94b8;font-size:0.88rem;">{len(_al_entries)} entr(ies)</p>', unsafe_allow_html=True)
    st.markdown("---")

    _outcome_colors = {
        "Answered":            "#22c55e",
        "Ticket Created":      "#ef4444",
        "Clarification Asked": "#f59e0b",
    }
    _intent_colors = {
        "question":      "#3b82f6",
        "generate":      "#22c55e",
        "compare":       "#9b7cf4",
        "clarification": "#f59e0b",
        "chitchat":      "#6b7280",
    }

    for _i, _entry in enumerate(_al_entries):
        _oc  = _outcome_colors.get(_entry.get("outcome", ""), "#6b7280")
        _ic  = _intent_colors.get(_entry.get("intent", ""), "#6b7280")
        _ts  = _entry.get("asked_at", "")
        _conf = _entry.get("confidence", 0)
        _conf_color = _cr_score_color(_conf) if _conf else "#4a6080"

        with st.expander(
            f"🤖 {_entry.get('question', '')[:80]}  —  {_ts}",
            expanded=(_i == 0)
        ):
            # Badges row
            _al_badges = (
                f'<span style="padding:2px 10px;border-radius:12px;font-size:0.68rem;font-weight:700;'                f'background:{_oc}22;color:{_oc};border:1px solid {_oc}55;margin-right:6px;">'                f'{_entry.get("outcome","")}</span>'                f'<span style="padding:2px 10px;border-radius:12px;font-size:0.68rem;font-weight:700;'                f'background:{_ic}22;color:{_ic};border:1px solid {_ic}55;margin-right:6px;">'                f'{_entry.get("intent","")}</span>'
            )
            if _conf:
                _al_badges += (
                    f'<span style="padding:2px 10px;border-radius:12px;font-size:0.68rem;font-weight:700;'                    f'background:{_conf_color}22;color:{_conf_color};border:1px solid {_conf_color}55;">'                    f'Confidence: {_conf:.0%}</span>'
                )
            st.markdown(_al_badges, unsafe_allow_html=True)
            st.markdown("")

            st.markdown("**Reply**")
            _reply_html = _entry.get("reply", "").replace("\n", "<br>")
            st.markdown(
                f'<div style="background:#0f1c2e;border:1px solid #1e3a5f;border-radius:8px;'                f'padding:14px 18px;margin:6px 0;color:#e2e8f0;line-height:1.7;font-size:0.85rem;">'                f'{_reply_html}</div>',
                unsafe_allow_html=True
            )

            if _entry.get("sources") and _entry["sources"] != "None":
                st.markdown("**Sources**")
                for _sl in _entry["sources"].split("\n"):
                    if _sl.strip():
                        st.markdown(
                            f'<div style="background:#111827;border:1px solid #1e3a5f;'                            f'border-left:3px solid #1a56db;border-radius:6px;'                            f'padding:6px 12px;margin-bottom:4px;font-size:0.8rem;color:#94a3b8;">'                            f'{_sl}</div>',
                            unsafe_allow_html=True
                        )

            if _entry.get("thread_id"):
                st.caption(f"Thread: {_entry['thread_id'][:8]}...")
            if _entry.get("url"):
                _al_url = _entry["url"]
                st.markdown(
                    f'<a href="{_al_url}" target="_blank" style="font-size:0.78rem;color:#60a5fa;">'                    f'Open in Notion ↗</a>',
                    unsafe_allow_html=True
                )
    st.stop()


# ─────────────────────────────────────────────────────────────────────────────
# PAGE: Ask Documents
# ─────────────────────────────────────────────────────────────────────────────

if page == "💬  Ask Documents":
    if not _cr_check_backend():
        st.error("⚠️ CiteRAG backend is not running. Start it with: `uvicorn backend.main:app --port 9000`")
        st.stop()

    st.markdown("""
    <div class="page-header">
        <h1>💬 Ask Documents</h1>
        <p>Ask any question — CiteRAG retrieves relevant chunks and answers with citations.</p>
    </div>
    """, unsafe_allow_html=True)

    doc_data = _cr_fetch_docs()

    col_q, col_f = st.columns([3, 1])
    with col_q:
        question = st.text_area("Your question",
            placeholder='e.g. "What are the key risks identified in the Project Brief?"',
            height=90, label_visibility="collapsed")
    with col_f:
        st.markdown("**Filters**")
        ask_industry = st.selectbox("Industry", ["All"] + doc_data.get("industries", []), key="ask_industry")
        ask_title    = st.selectbox("Document",  ["All"] + doc_data.get("titles", []),    key="ask_title")
        ask_topk     = st.slider("Sources", 3, 10, 6, key="ask_topk")

    if st.button("⚡ Ask", use_container_width=True, type="primary"):
        if not question.strip():
            st.warning("Please enter a question.")
        else:
            filters = {}
            if ask_industry != "All": filters["industry"]  = ask_industry
            if ask_title    != "All": filters["doc_title"] = ask_title
            with st.spinner("Searching and generating answer..."):
                res = requests.post(f"{CITERAG_BACKEND}/ask", json={
                    "question": question, "top_k": ask_topk,
                    "filters":  filters or None,
                })
            if res.status_code == 200:
                st.session_state["cr_last_ask"] = {
                    "question": question,
                    "answer":   res.json()["answer"],
                    "sources":  res.json()["sources"],
                }
            else:
                st.error(f"Error: {res.text[:200]}")

    if "cr_last_ask" in st.session_state:
        last = st.session_state["cr_last_ask"]
        answer_html = last["answer"].replace("\n", "<br>")
        st.markdown(f"""
        <div style="background:#0f1c2e;border:1px solid #1e3a5f;border-radius:8px;
                    padding:20px 24px;margin:12px 0;color:#e2e8f0;line-height:1.7;font-size:0.92rem;">
            {answer_html}
        </div>""", unsafe_allow_html=True)
        _cr_render_sources(last["sources"])

    st.stop()


# ─────────────────────────────────────────────────────────────────────────────
# PAGE: Refine Document
# ─────────────────────────────────────────────────────────────────────────────

if page == "🖊️  Refine Document":
    st.markdown("""
    <div class="page-header">
        <h1>🖊️ Refine Document</h1>
        <p>Load any saved document, regenerate it entirely or refine specific sections, then download, push to Notion or save.</p>
    </div>
    """, unsafe_allow_html=True)

    # ── Step 1: Pick a document from DB ──────────────────────────────────────
    st.markdown('<div class="step-label">Step 1 — Select a Document</div>', unsafe_allow_html=True)

    with st.spinner("Loading history..."):
        _rf_res = requests.get(f"{BACKEND}/documents")
    if _rf_res.status_code != 200:
        st.error("Could not load document history. Make sure the backend is running.")
        st.stop()

    _rf_docs = _rf_res.json().get("documents", [])
    if not _rf_docs:
        st.info("No documents saved yet. Generate and save a document first.")
        st.stop()

    # Build label → doc mapping
    _rf_labels = {
        f"{d['title']} — {d.get('doc_format','word').upper()} v{d.get('version',1)} (ID #{d['id']})": d
        for d in _rf_docs
    }
    _rf_choice = st.selectbox("Choose document", list(_rf_labels.keys()),
                               label_visibility="collapsed", key="rf_doc_choice")
    _rf_doc    = _rf_labels[_rf_choice]
    _rf_id     = _rf_doc["id"]
    _rf_fmt    = _rf_doc.get("doc_format", "word")
    _rf_title  = _rf_doc["title"]
    _rf_dtype  = _rf_doc.get("doc_type", "General")

    if st.button("📂 Load Document", use_container_width=True, key="rf_load_btn", type="primary"):
        with st.spinner("Loading..."):
            _rf_detail = requests.get(f"{BACKEND}/documents/{_rf_id}")
        if _rf_detail.status_code == 200:
            _rf_data = _rf_detail.json()
            st.session_state["rf_content"]   = _rf_data.get("content", {})
            st.session_state["rf_doc_id"]    = _rf_id
            st.session_state["rf_title"]     = _rf_title
            st.session_state["rf_format"]    = _rf_fmt
            st.session_state["rf_doc_type"]  = _rf_dtype
            st.session_state.pop("rf_refined_buf", None)
            st.rerun()
        else:
            st.error("Could not load document.")

    # ── Step 2: Show + Refine ─────────────────────────────────────────────────
    if "rf_content" in st.session_state and st.session_state.get("rf_doc_id") == _rf_id:
        rf_content = st.session_state["rf_content"]
        rf_fmt     = st.session_state["rf_format"]
        rf_title   = st.session_state["rf_title"]
        rf_dtype   = st.session_state["rf_doc_type"]

        st.divider()
        st.markdown('<div class="step-label">Step 2 — Preview</div>', unsafe_allow_html=True)

        if rf_fmt == "excel":
            render_excel_document(rf_title, rf_content.get("sheets", []))
        else:
            render_full_document(rf_title, rf_content)

        # ── Regenerate entire document ────────────────────────────────────────
        st.divider()
        st.markdown('<div class="step-label">Step 3 — Refine</div>', unsafe_allow_html=True)

        with st.expander("🔄 Regenerate entire document", expanded=False):
            st.caption("Describe what you want changed overall — the LLM will rewrite the whole document.")
            rf_regen_fb = st.text_area("What should change overall?",
                placeholder="e.g. Make it more formal • Add an executive summary • Change the tone to persuasive",
                key="rf_regen_feedback", height=90)
            if st.button("🔄 Regenerate Document", key="rf_regen_btn", use_container_width=True):
                if not rf_regen_fb.strip():
                    st.warning("Please describe what should change.")
                else:
                    # Build sections list from content
                    _meta = {"__section_order__", "__show_headings__"}
                    _ord  = rf_content.get("__section_order__")
                    if _ord:
                        _sections = [k for k in _ord if k in rf_content and k not in _meta]
                    else:
                        _sections = [k for k in rf_content if k not in _meta]

                    with st.spinner("Regenerating..."):
                        _rg_res = requests.post(f"{BACKEND}/generate", json={
                            "title":    rf_title,
                            "sections": _sections,
                            "answers":  {f"Overall feedback for {rf_title}": rf_regen_fb},
                            "doc_format": rf_fmt,
                            "show_headings": rf_content.get("__show_headings__", True),
                        })
                    if _rg_res.status_code == 200:
                        _new = _rg_res.json()
                        # Preserve meta flags from original
                        _new["__show_headings__"] = rf_content.get("__show_headings__", True)
                        _new["__section_order__"] = _sections
                        st.session_state["rf_content"] = _new
                        st.session_state.pop("rf_refined_buf", None)
                        st.success("Document regenerated!")
                        st.rerun()
                    else:
                        st.error(f"Regeneration failed: {_rg_res.text[:200]}")

        # ── Refine individual sections ────────────────────────────────────────
        if rf_fmt == "excel":
            sheets = rf_content.get("sheets", [])
            for _si, _sheet in enumerate(sheets):
                _sname = _sheet.get("sheet_name", f"Sheet {_si+1}")
                with st.expander(f"✏️ Refine: {_sname}", expanded=False):
                    _headers = _sheet.get("headers", [])
                    _rows    = _sheet.get("rows", [])
                    if _headers and _rows:
                        _mc = max(len(_headers), max((len(r) for r in _rows), default=0))
                        _h2 = _headers + [f"Col {j+1}" for j in range(len(_headers), _mc)]
                        _pd = [r + [""] * max(0, len(_h2)-len(r)) for r in _rows]
                        st.dataframe(pd.DataFrame(_pd, columns=_h2), use_container_width=True, hide_index=True)
                    _sfb = st.text_area(f"Changes for '{_sname}'",
                        placeholder="e.g. Update total to 500000 • Add row: Tax = 45000",
                        key=f"rf_sheet_fb_{_si}", height=80)
                    if st.button(f"Update '{_sname}'", key=f"rf_sheet_btn_{_si}", use_container_width=True):
                        with st.spinner(f"Refining {_sname}..."):
                            _sr = requests.post(f"{BACKEND}/refine-section", json={
                                "section_name": _sname, "current_data": _sheet,
                                "feedback": _sfb, "doc_format": "excel"
                            })
                        if _sr.status_code == 200:
                            _us = _sr.json().get("updated_sheet", _sheet)
                            if _us.get("rows"):
                                _new_sheets = list(sheets)
                                _new_sheets[_si] = _us
                                st.session_state["rf_content"] = {**rf_content, "sheets": _new_sheets}
                                st.session_state.pop("rf_refined_buf", None)
                                st.success(f"'{_sname}' updated!")
                                st.rerun()
                            else:
                                st.error("Refinement returned empty — try rephrasing.")
                        elif _sr.status_code == 429:
                            st.warning("⏳ Too many refinements — please wait a moment.")
                        else:
                            st.error(f"Error: {_sr.text[:200]}")
        else:
            _meta = {"__section_order__", "__show_headings__"}
            _ord  = rf_content.get("__section_order__")
            _sec_items = [(k, rf_content[k]) for k in _ord if k in rf_content] if _ord                          else [(k, v) for k, v in rf_content.items() if k not in _meta]
            for _sec, _content in _sec_items:
                with st.expander(f"✏️ Refine: {_sec}", expanded=False):
                    st.write(flatten_to_text(_content)[:600] + ("…" if len(flatten_to_text(_content)) > 600 else ""))
                    _sfb2 = st.text_area(f"Changes for '{_sec}'",
                        placeholder="e.g. Make it more concise • Add a specific example • Change tone",
                        key=f"rf_sec_fb_{_sec}", height=80)
                    if st.button(f"Update '{_sec}'", key=f"rf_sec_btn_{_sec}", use_container_width=True):
                        with st.spinner(f"Refining '{_sec}'..."):
                            _sr2 = requests.post(f"{BACKEND}/refine-section", json={
                                "section_name": _sec,
                                "original_text": flatten_to_text(_content),
                                "feedback": _sfb2, "doc_format": "word"
                            })
                        if _sr2.status_code == 200:
                            _ut = _sr2.json().get("updated_text", "")
                            if isinstance(_ut, (dict, list)):
                                _ut = flatten_to_text(_ut)
                            st.session_state["rf_content"] = {**rf_content, _sec: str(_ut)}
                            st.session_state.pop("rf_refined_buf", None)
                            st.success(f"'{_sec}' updated!")
                            st.rerun()
                        elif _sr2.status_code == 429:
                            st.warning("⏳ Too many refinements — please wait a moment.")
                        else:
                            st.error(f"Error: {_sr2.text[:200]}")

        # ── Actions card ──────────────────────────────────────────────────────
        st.divider()
        st.markdown('<div class="step-label">Actions</div>', unsafe_allow_html=True)

        if rf_fmt == "excel":
            try:
                _rf_buf = create_excel_file_from_data(rf_content)
            except Exception:
                _rf_buf = None
        else:
            _rf_buf = create_word_document(rf_title, rf_content)

        _render_doc_card(
            doc={
                "doc_id":     _rf_id,
                "title":      rf_title,
                "doc_format": rf_fmt,
                "doc_type":   rf_dtype,
                "file_ext":   "xlsx" if rf_fmt == "excel" else "docx",
                "version":    _rf_doc.get("version", 1),
                "created_at": _rf_doc.get("created_at", ""),
            },
            in_memory_content = rf_content,
            in_memory_buf     = _rf_buf,
        )

    st.stop()


# ─────────────────────────────────────────────────────────────────────────────
# PAGE: Compare
# ─────────────────────────────────────────────────────────────────────────────

if page == "⚖️  Compare":
    if not _cr_check_backend():
        st.error("⚠️ CiteRAG backend is not running.")
        st.stop()

    st.markdown("""
    <div class="page-header">
        <h1>⚖️ Compare Documents</h1>
        <p>Pick two documents and a focus area — get a structured side-by-side comparison.</p>
    </div>
    """, unsafe_allow_html=True)

    titles = _cr_fetch_docs().get("titles", [])
    if len(titles) < 2:
        st.info("At least 2 documents need to be indexed to use compare mode. Go to 📥 Ingest first.")
        st.stop()

    c_col1, c_col2 = st.columns(2)
    with c_col1: doc_a = st.selectbox("Document A", titles, key="cmp_a")
    with c_col2: doc_b = st.selectbox("Document B", [t for t in titles if t != doc_a], key="cmp_b")

    focus = st.text_input("Comparison focus", value="overall content and structure",
        placeholder='e.g. "risk sections" or "financial assumptions"')

    if st.button("⚖️ Compare", use_container_width=True, type="primary"):
        if doc_a == doc_b:
            st.warning("Please select two different documents.")
        else:
            with st.spinner("Comparing documents..."):
                res = requests.post(f"{CITERAG_BACKEND}/compare", json={
                    "title_a": doc_a, "title_b": doc_b, "focus": focus,
                })
            if res.status_code == 200:
                data   = res.json()
                src_c1, src_c2 = st.columns(2)
                with src_c1:
                    st.markdown(f"**📄 {doc_a}** — retrieved chunks")
                    _cr_render_sources(data.get("doc_a_chunks", []), expanded=False)
                with src_c2:
                    st.markdown(f"**📄 {doc_b}** — retrieved chunks")
                    _cr_render_sources(data.get("doc_b_chunks", []), expanded=False)
                st.markdown("---")
                st.markdown("### Comparison")
                st.markdown(data.get("comparison", ""))
            else:
                st.error(f"Compare error: {res.text[:200]}")
    st.stop()


# ─────────────────────────────────────────────────────────────────────────────
# PAGE: Eval Lab
# ─────────────────────────────────────────────────────────────────────────────

if page == "🟰  Evaluate":
    if not _cr_check_backend():
        st.error("⚠️ CiteRAG backend is not running.")
        st.stop()

    st.markdown("""
    <div class="page-header">
        <h1>🟰 Evaluate</h1>
        <p>Run RAGAS evaluation to measure retrieval and answer quality.</p>
    </div>
    """, unsafe_allow_html=True)

    doc_data_e = _cr_fetch_docs()
    e_col1, e_col2 = st.columns([2, 1])

    with e_col1:
        run_name      = st.text_input("Run name", placeholder="e.g. baseline-top5")
        questions_raw = st.text_area("Evaluation questions (one per line)", height=180,
            placeholder="What are the key deliverables?\nWhat budget was allocated?\nWhat risks were identified?")
    with e_col2:
        e_topk     = st.slider("top_k", 3, 10, 5, key="eval_topk")
        e_industry = st.selectbox("Industry filter", ["None"] + doc_data_e.get("industries", []))
        e_title    = st.selectbox("Doc filter",      ["None"] + doc_data_e.get("titles", []))
        st.info("Measures:\n- Faithfulness\n- Answer relevancy")

    if st.button("▶ Run Evaluation", use_container_width=True, type="primary"):
        questions = [q.strip() for q in questions_raw.strip().split("\n") if q.strip()]
        if not questions:       st.warning("Please enter at least one question.")
        elif not run_name.strip(): st.warning("Please give this run a name.")
        else:
            filters = {}
            if e_industry != "None": filters["industry"]  = e_industry
            if e_title    != "None": filters["doc_title"] = e_title
            with st.spinner(f"Running RAGAS on {len(questions)} question(s)..."):
                res = requests.post(f"{CITERAG_BACKEND}/eval/run", json={
                    "questions": questions,
                    "config": {"run_name": run_name, "top_k": e_topk, "filters": filters or None},
                })
            if res.status_code == 200:
                data    = res.json()
                summary = data.get("summary", {})
                st.success(f"✅ Eval run saved (ID #{data.get('run_id')})")
                cols = st.columns(len(summary))
                for col, (metric, score) in zip(cols, summary.items()):
                    color = _cr_score_color(score)
                    col.markdown(f"""
                    <div style="background:#111827;border:1px solid #1e2d40;border-radius:6px;
                                padding:12px 16px;text-align:center;">
                        <div style="font-size:0.72rem;font-weight:600;letter-spacing:0.1em;
                                    color:#64748b;text-transform:uppercase;margin-bottom:4px;">
                            {metric.replace('_',' ').title()}
                        </div>
                        <div style="font-size:1.8rem;font-weight:700;color:{color};">{score:.0%}</div>
                    </div>""", unsafe_allow_html=True)
                with st.expander("📋 Per-question results", expanded=False):
                    for r in data.get("results", []):
                        st.markdown(f"**Q:** {r['question']}")
                        st.markdown(f"**A:** {r['answer'][:400]}...")
                        st.markdown("---")
            elif res.status_code == 422:
                st.error("RAGAS not installed in CiteRAG. Run: `pip install ragas datasets`")
            else:
                st.error(f"Eval error: {res.text[:200]}")

    st.markdown("---")
    st.markdown("### Past Runs")
    try:
        runs_res = requests.get(f"{CITERAG_BACKEND}/eval/runs", timeout=5)
        if runs_res.status_code == 200:
            runs = runs_res.json().get("runs", [])
            if not runs:
                st.caption("No evaluation runs yet.")
            else:
                for run in runs:
                    summary = run.get("summary", {})
                    scores  = "  ·  ".join(
                        f"{k.replace('_',' ').title()}: **{v:.0%}**" for k, v in summary.items())
                    with st.expander(f"#{run['id']} — {run['run_name']}  ({run['created_at']})"):
                        st.markdown(scores)
                        cfg = run.get("config", {})
                        st.caption(f"top_k={cfg.get('top_k')}  ·  filters={cfg.get('filters')}")
    except Exception:
        st.caption("Could not load past runs.")
    st.stop()


# ─────────────────────────────────────────────────────────────────────────────
# PAGE: Ingest
# ─────────────────────────────────────────────────────────────────────────────
# PAGE: CiteRAG Log
# ─────────────────────────────────────────────────────────────────────────────

if page == "📋  CiteRAG Log":
    if not _cr_check_backend():
        st.error("⚠️ CiteRAG backend is not running.")
        st.stop()

    st.markdown("""
    <div class="page-header">
        <h1>📋 CiteRAG Log</h1>
        <p>All Ask, Compare and Eval interactions are automatically saved here and in Notion.</p>
    </div>
    """, unsafe_allow_html=True)

    log_limit = st.slider("Entries to load", 10, 100, 30, key="log_limit")
    if st.button("🔄 Refresh", use_container_width=False):
        st.rerun()

    try:
        log_res = requests.get(f"{CITERAG_BACKEND}/qa-log", params={"limit": log_limit}, timeout=10)
        if log_res.status_code == 200:
            entries = log_res.json().get("entries", [])
            if not entries:
                st.info("No entries yet. Ask a question (💬), run a comparison (⚖️) or an eval (🧪) to see logs here.")
            else:
                st.markdown(f"Showing **{len(entries)}** most recent entries")
                for i, entry in enumerate(entries):
                    conf  = entry.get("confidence", 0)
                    color = _cr_score_color(conf)
                    asked = entry.get("asked_at", "")[:16].replace("T", " ")
                    url   = entry.get("url", "")
                    notion_link = (
                        f'&nbsp;<a href="{url}" target="_blank" '
                        f'style="font-size:0.7rem;color:#60a5fa;">Open in Notion ↗</a>'
                    ) if url else ""

                    # Type badge styling
                    _etype = entry.get("type", "Ask")
                    _type_styles = {
                        "Ask":       ("#3b82f6", "💬"),
                        "Compare":   ("#9b7cf4", "⚖️"),
                        "Eval":      ("#f5c842", "🧪"),
                        "Assistant": ("#22c55e", "🤖"),
                    }
                    _tc, _te = _type_styles.get(_etype, ("#3b82f6", "💬"))

                    with st.expander(
                        f"{_te} **{entry.get('question', '')[:90]}**  —  {asked}",
                        expanded=(i == 0)
                    ):
                        st.markdown(
                            f'<div style="display:flex;align-items:center;gap:8px;margin-bottom:8px;">' +
                            f'<span style="padding:2px 10px;border-radius:12px;font-size:0.68rem;font-weight:700;' +
                            f'background:{_tc}22;color:{_tc};border:1px solid {_tc}55;">{_etype}</span>' +
                            f'<span style="display:inline-block;padding:2px 8px;border-radius:4px;' +
                            f'font-size:0.7rem;font-weight:700;background:{color}22;color:{color};' +
                            f'border:1px solid {color}55;">Confidence: {conf:.0%}</span>{notion_link}' +
                            f'</div>',
                            unsafe_allow_html=True
                        )
                        st.markdown("")
                        st.markdown("**Answer**")
                        answer_html = entry.get("answer", "").replace("\n", "<br>")
                        st.markdown(
                            f'<div style="background:#0f1c2e;border:1px solid #1e3a5f;border-radius:8px;'
                            f'padding:16px 20px;margin:8px 0;color:#e2e8f0;line-height:1.7;font-size:0.85rem;">'
                            f'{answer_html}</div>',
                            unsafe_allow_html=True
                        )
                        if entry.get("sources"):
                            st.markdown("**Sources**")
                            for line in entry["sources"].split("\n"):
                                if line.strip():
                                    st.markdown(
                                        f'<div style="background:#111827;border:1px solid #1e3a5f;'
                                        f'border-left:3px solid #1a56db;border-radius:6px;'
                                        f'padding:8px 14px;margin-bottom:6px;font-size:0.82rem;color:#94a3b8;">'
                                        f'{line}</div>',
                                        unsafe_allow_html=True
                                    )
                        if entry.get("filters") and entry["filters"] != "None":
                            st.caption(f"Filters: {entry['filters']}")
        else:
            st.error(f"Could not load CiteRAG log: {log_res.text[:200]}")
    except Exception as e:
        st.error(f"Error loading log: {e}")
    st.stop()


# ─────────────────────────────────────────────────────────────────────────────
# Separator divider in nav — skip if selected
# ─────────────────────────────────────────────────────────────────────────────
if page == "─────────────────":
    st.info("Please select a page from the sidebar.")
    st.stop()

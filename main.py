"""
main.py — DocForgeHub unified backend
======================================
Combines DocForge (document generation) and CiteRAG (RAG Q&A) into
a single FastAPI application running on one port.

DocForge routes  — /plan  /questions  /generate  /refine-section
                   /export/excel  /documents/*  /notion/*
                   /redis/status  /health

CiteRAG routes   — /rag/ingest  /rag/ingest/status  /rag/ingest/{page_id}
                   /rag/search  /rag/ask  /rag/ask/refine  /rag/qa-log
                   /rag/compare  /rag/documents
                   /rag/eval/run  /rag/eval/runs  /rag/eval/runs/{id}
                   /rag/health
"""

import io
import logging
import threading
import time
from fastapi import FastAPI, HTTPException, APIRouter
from fastapi.responses import StreamingResponse, JSONResponse, Response
from fastapi.middleware.cors import CORSMiddleware
from contextlib import asynccontextmanager

# ── DocForge imports ──────────────────────────────────────────────────────────
from backend.services.p1.planner import plan_document
from backend.services.p1.question import generate_questions
from backend.services.p1.generator import generate_document_sections
from backend.services.p1.excel_generator import generate_excel_sections, refine_excel_section
from backend.services.p1.refinement import refine_section
from backend.services.p1.excel_exporter import generate_excel_file
from backend.services.p1.notion import push_to_notion, update_notion_page
from backend.database import (
    init_db, save_document, list_documents, get_document,
    delete_document, list_versions_by_id, delete_all_versions, get_latest_version
)
from backend.services.p1.redis import redis_svc, ThrottleExceeded

# ── CiteRAG imports ───────────────────────────────────────────────────────────
from backend.chroma_db import (
    init_db as init_chroma_db,
    get_all_titles, get_all_industries, list_ingested_docs,
    delete_chunks_by_page, list_eval_runs, get_eval_run,
)

logging.basicConfig(level=logging.INFO)
log = logging.getLogger(__name__)

# ── Assistant imports ─────────────────────────────────────────────────────────
from backend.services.p3.state             import init_assistant_table, list_threads
from backend.services.p3.graph             import run_turn
from backend.services.p3.tickets    import fetch_tickets, update_ticket_status
from backend.services.p3.assistant_log import fetch_assistant_log

# ── Ingest config ─────────────────────────────────────────────────────────────
INGEST_POLL_INTERVAL = 15 * 60   # check Notion for new/updated docs every 15 min


def _auto_ingest(page_ids: list = None, force: bool = False):
    """
    Ingest Notion docs into ChromaDB.

    - On startup:         runs for all pages, skips unchanged ones (force=False)
    - After Notion push:  force=True on the specific page_id just pushed
    - Periodic poll:      runs every INGEST_POLL_INTERVAL seconds, skips unchanged

    The skip logic in ingest_notion() compares Notion last_edited_time
    against stored ingested_at — so unchanged docs cost only a metadata fetch.
    """
    try:
        from backend.services.p2.ingestion import ingest_notion
        if page_ids:
            log.info(f"Auto-ingest: syncing page(s) {page_ids}...")
        else:
            log.info("Auto-ingest: checking for new/updated Notion docs...")
        result   = ingest_notion(page_ids=page_ids, force=force)
        ingested = result.get("ingested", 0)
        skipped  = result.get("skipped", 0)
        chunks   = result.get("chunks", 0)
        if ingested:
            log.info(
                f"Auto-ingest: {ingested} new/updated doc(s) → {chunks} chunks "
                f"({skipped} already up to date)"
            )
        else:
            log.info(f"Auto-ingest: all {skipped} doc(s) already up to date — nothing to do")
    except Exception as e:
        log.warning(f"Auto-ingest failed (non-fatal): {e}")


def _ingest_poll_loop():
    """
    Background daemon thread: polls Notion every INGEST_POLL_INTERVAL seconds.
    Skips docs that haven't changed since last ingest.
    Runs forever until the process exits.
    """
    # Wait one full interval before first poll (startup already ran ingest)
    time.sleep(INGEST_POLL_INTERVAL)
    while True:
        log.info("Auto-ingest poll: checking Notion for new/updated docs...")
        _auto_ingest()
        time.sleep(INGEST_POLL_INTERVAL)


# ── Lifespan — init both databases on startup ─────────────────────────────────
@asynccontextmanager
async def lifespan(app: FastAPI):
    # DocForge — PostgreSQL
    init_db()
    log.info("DocForge PostgreSQL ready")
    # DocForge — Redis
    status = redis_svc.status()
    log.info(status["message"])
    # CiteRAG — ChromaDB + SQLite
    init_chroma_db()
    log.info("CiteRAG ChromaDB + SQLite ready")
    # Assistant — thread state table
    try:
        init_assistant_table()
        log.info("Assistant thread table ready")
    except Exception as _e:
        log.error(f"init_assistant_table failed: {_e}")
    # Startup ingest — skips docs already indexed and unchanged
    threading.Thread(target=_auto_ingest, daemon=True).start()
    # Periodic poll — re-checks every INGEST_POLL_INTERVAL minutes
    threading.Thread(target=_ingest_poll_loop, daemon=True).start()
    log.info(f"Auto-ingest started (polling every {INGEST_POLL_INTERVAL // 60} min)")
    yield


app = FastAPI(
    title       = "DocForgeHub API",
    description = "Unified backend: DocForge document generation + CiteRAG RAG Q&A",
    version     = "2.0.0",
    lifespan    = lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ══════════════════════════════════════════════════════════════════════════════
# DOCFORGE ROUTES
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/plan")
def plan(payload: dict):
    prompt = payload["prompt"]
    cached = redis_svc.get_cached_plan(prompt)
    if cached:
        return cached
    result = plan_document(prompt)
    redis_svc.cache_plan(prompt, result)
    return result


@app.post("/questions")
def questions(payload: dict):
    title    = payload["title"]
    sections = payload["sections"]
    cached   = redis_svc.get_cached_questions(title, sections)
    if cached:
        return cached
    result = generate_questions(title, sections)
    redis_svc.cache_questions(title, sections, result)
    return result


@app.post("/generate")
def generate(payload: dict):
    doc_format = payload.get("doc_format", "word")
    cached     = redis_svc.get_cached_generation(payload["title"], payload["sections"], doc_format)
    if cached:
        return cached
    show_headings = payload.get("show_headings", True)
    if doc_format == "excel":
        result = generate_excel_sections(
            payload["title"], payload["sections"], payload["answers"]
        )
    else:
        result = generate_document_sections(
            payload["title"], payload["sections"], payload["answers"],
            show_headings=show_headings,
        )
        if isinstance(result, dict):
            result["__show_headings__"] = show_headings
    redis_svc.cache_generation(payload["title"], payload["sections"], doc_format, result)
    return result


@app.post("/refine-section")
def refine(payload: dict):
    title = payload.get("section_name", "unknown")
    try:
        redis_svc.check_refine_limit(title)
    except ThrottleExceeded as e:
        raise HTTPException(status_code=429, detail=f"Too many refinements. ({e})")
    doc_format = payload.get("doc_format", "word")
    if doc_format == "excel":
        updated_sheet = refine_excel_section(
            payload["section_name"], payload.get("current_data", {}), payload["feedback"]
        )
        return {"updated_sheet": updated_sheet}
    else:
        updated = refine_section(
            payload["section_name"], payload["original_text"], payload["feedback"]
        )
        return {"updated_text": updated}


@app.post("/export/excel")
def export_excel(payload: dict):
    try:
        buf = generate_excel_file(payload)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=document.xlsx"}
        )
    except Exception as e:
        return JSONResponse(status_code=500, content={"detail": str(e)})


@app.post("/documents/save")
def save_doc(payload: dict):
    try:
        raw_bytes    = bytes.fromhex(payload["file_bytes"]) if payload.get("file_bytes") else None
        save_mode    = payload.get("save_mode", "new_version")
        overwrite_id = payload.get("overwrite_id")
        if save_mode == "overwrite" and not overwrite_id:
            raise HTTPException(status_code=400, detail="overwrite_id required for overwrite mode")
        result = save_document(
            title        = payload["title"],
            doc_type     = payload.get("doc_type", "General"),
            doc_format   = payload.get("doc_format", "word"),
            content      = payload.get("content", {}),
            file_bytes   = raw_bytes,
            file_ext     = payload.get("file_ext"),
            save_mode    = save_mode,
            overwrite_id = overwrite_id,
        )
        mode_label = "overwritten" if result["mode"] == "overwrite" else f"saved as v{result['version']}"
        return {**result, "message": f"Document {mode_label}"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/documents")
def get_documents():
    try:
        docs = list_documents()
        for doc in docs:
            if doc.get("created_at"):
                doc["created_at"] = doc["created_at"].strftime("%d %b %Y, %I:%M %p")
            doc["version_label"] = f"v{doc.get('version', 1)}"
        return {"documents": docs}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/documents/{doc_id}")
def get_doc(doc_id: int):
    doc = get_document(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    doc.pop("file_bytes", None)
    if doc.get("created_at"):
        doc["created_at"] = doc["created_at"].strftime("%d %b %Y, %I:%M %p")
    return doc


@app.get("/documents/{doc_id}/download")
def download_doc(doc_id: int):
    doc = get_document(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    ext        = doc.get("file_ext") or ("xlsx" if doc.get("doc_format") == "excel" else "docx")
    safe_title = doc["title"].replace(" ", "_")[:40]
    mime = (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        if ext == "xlsx"
        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    if doc.get("file_bytes"):
        return Response(
            content=doc["file_bytes"], media_type=mime,
            headers={"Content-Disposition": f"attachment; filename={safe_title}.{ext}"}
        )
    content = doc.get("content", {})
    try:
        if ext == "xlsx":
            file_bytes = generate_excel_file(content)
            if hasattr(file_bytes, "read"):
                file_bytes = file_bytes.read()
        else:
            from docx import Document as DocxDocument
            from docx.shared import Pt
            from io import BytesIO
            d = DocxDocument()
            d.add_heading(doc["title"], 0).alignment = 1
            order     = content.get("__section_order__")
            meta      = {"__section_order__", "__show_headings__"}
            show_hdgs = content.get("__show_headings__", True)
            if order:
                items   = [(k, content[k]) for k in order if k in content]
                covered = set(order) | meta
                items  += [(k, v) for k, v in content.items() if k not in covered]
            else:
                items = [(k, v) for k, v in content.items() if k not in meta]
            def _flat(v):
                if v is None: return ""
                if isinstance(v, str): return v.strip()
                if isinstance(v, list): return "\n".join(f"• {_flat(i)}" for i in v)
                if isinstance(v, dict):
                    return "\n\n".join(
                        f"{str(k).replace('_',' ').title()}:\n{_flat(val)}"
                        for k, val in v.items() if k not in meta
                    )
                return str(v)

            # Add TOC page for multi-section docs with headings
            if show_hdgs and len(items) > 1:
                from docx.shared import RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                toc_h = d.add_paragraph()
                toc_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
                tr = toc_h.add_run("Table of Contents")
                tr.bold = True
                tr.font.size = Pt(14)
                tr.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
                toc_h.paragraph_format.space_after = Pt(6)
                for idx, (sn, _) in enumerate(items, 1):
                    label = str(sn).replace("_", " ").title()
                    tp = d.add_paragraph()
                    tp.paragraph_format.space_before = Pt(2)
                    tp.paragraph_format.space_after  = Pt(2)
                    rl = tp.add_run(f"{idx}.  {label}")
                    rl.font.size = Pt(11)
                    rl.font.color.rgb = RGBColor(0x1E, 0x3A, 0x6E)
                    rd = tp.add_run(f"  {'.' * max(2, 50 - len(label))}")
                    rd.font.size = Pt(9)
                    rd.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
                d.add_page_break()

            # Add TOC page for multi-section docs with headings
            if show_hdgs and len(items) > 1:
                from docx.shared import RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml.ns import qn as _qn
                from docx.oxml import OxmlElement as _OE
                toc_p = d.add_paragraph()
                toc_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _r = toc_p.add_run()
                _fb = _OE("w:fldChar"); _fb.set(_qn("w:fldCharType"), "begin"); _r._r.append(_fb)
                _it = _OE("w:instrText"); _it.set(_qn("xml:space"), "preserve")
                _it.text = ' TOC \\o "1-3" \\h \\z \\u '; _r._r.append(_it)
                _fs = _OE("w:fldChar"); _fs.set(_qn("w:fldCharType"), "separate"); _r._r.append(_fs)
                _fe = _OE("w:fldChar"); _fe.set(_qn("w:fldCharType"), "end"); _r._r.append(_fe)
                _np = d.add_paragraph()
                _nr = _np.add_run("(Open in Microsoft Word and press Ctrl+A then F9 to update page numbers)")
                _nr.font.size = Pt(9); _nr.font.italic = True
                _nr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
                _np.paragraph_format.space_after = Pt(8)
                for _idx, (_sn, _) in enumerate(items, 1):
                    _lbl = str(_sn).replace("_", " ").title()
                    _tp = d.add_paragraph()
                    _tp.paragraph_format.space_before = Pt(2)
                    _tp.paragraph_format.space_after  = Pt(2)
                    _rl = _tp.add_run(f"{_idx}.  {_lbl}")
                    _rl.font.size = Pt(11); _rl.font.color.rgb = RGBColor(0x1E, 0x3A, 0x6E)
                    _rd = _tp.add_run(f"  {'.' * max(2, 50 - len(_lbl))}")
                    _rd.font.size = Pt(9); _rd.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
                d.add_page_break()

            for section_name, section_content in items:
                if show_hdgs:
                    d.add_heading(str(section_name).replace("_", " ").title(), level=1)
                for para in [p.strip() for p in _flat(section_content).split("\n") if p.strip()]:
                    if para.startswith("•"):
                        d.add_paragraph(para.lstrip("• ").strip(), style="List Bullet")
                    else:
                        p = d.add_paragraph()
                        p.add_run(para).font.size = Pt(11)
                d.add_paragraph()
            buf = BytesIO()
            d.save(buf)
            buf.seek(0)
            file_bytes = buf.read()
        return Response(
            content=file_bytes, media_type=mime,
            headers={"Content-Disposition": f"attachment; filename={safe_title}.{ext}"}
        )
    except Exception as e:
        log.error(f"Failed to regenerate file for doc {doc_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Could not regenerate file: {e}")


@app.patch("/documents/{doc_id}/rename")
def rename_doc(doc_id: int, payload: dict):
    """Rename a document (updates all versions with same title)."""
    new_title = (payload.get("title") or "").strip()
    if not new_title:
        raise HTTPException(status_code=400, detail="Title cannot be empty")
    doc = get_document(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    try:
        from backend.database import get_conn, release_conn
        conn = get_conn()
        old_title = doc["title"]
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "UPDATE documents SET title = %s WHERE title = %s",
                    (new_title, old_title)
                )
                count = cur.rowcount
            conn.commit()
        finally:
            release_conn(conn)
        return {"message": f"Renamed to '{new_title}' ({count} version(s) updated)"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/documents/{doc_id}")
def delete_doc(doc_id: int):
    if not delete_document(doc_id):
        raise HTTPException(status_code=404, detail="Document not found")
    return {"message": "Document deleted"}


@app.post("/documents/check-version")
def check_version(payload: dict):
    latest = get_latest_version(payload.get("title", ""))
    if not latest:
        return {"exists": False}
    return {"exists": True, "latest_id": latest["id"], "latest_version": latest["version"]}


@app.get("/documents/{doc_id}/versions")
def get_doc_versions(doc_id: int):
    versions = list_versions_by_id(doc_id)
    if not versions:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"versions": versions, "total": len(versions)}


@app.delete("/documents/{doc_id}/all-versions")
def delete_doc_all_versions(doc_id: int):
    doc = get_document(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    count = delete_all_versions(doc["title"])
    return {"message": f"Deleted {count} version(s) of '{doc['title']}'"}


@app.post("/notion/push")
def notion_push(payload: dict):
    try:
        db_id    = payload.get("db_id")
        version  = 1
        doc_type = payload.get("doc_type", "General")
        if db_id:
            doc = get_document(db_id)
            if doc:
                version  = doc.get("version", 1)
                if not payload.get("doc_type"):
                    doc_type = doc.get("doc_type", "General")
        result = push_to_notion(
            title      = payload["title"],
            doc_format = payload.get("doc_format", "word"),
            content    = payload.get("content", {}),
            db_id      = db_id,
            version    = version,
            doc_type   = doc_type,
        )
        # Force re-ingest the specific page just pushed
        _new_page_id = result.get("page_id")
        threading.Thread(
            target=_auto_ingest,
            kwargs={"page_ids": [_new_page_id] if _new_page_id else None, "force": True},
            daemon=True,
        ).start()
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Notion push failed: {e}")


@app.post("/notion/update")
def notion_update(payload: dict):
    try:
        result = update_notion_page(
            page_id    = payload["page_id"],
            title      = payload["title"],
            doc_format = payload.get("doc_format", "word"),
            content    = payload.get("content", {}),
            version    = payload.get("version", 1),
            doc_type   = payload.get("doc_type", "General"),
        )
        # Force re-ingest the updated page
        threading.Thread(
            target=_auto_ingest,
            kwargs={"page_ids": [payload["page_id"]], "force": True},
            daemon=True,
        ).start()
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Notion update failed: {e}")


@app.get("/redis/status")
def redis_status():
    return redis_svc.status()


@app.get("/health")
def health():
    return {"status": "ok", "service": "DocForgeHub", "redis": redis_svc.is_available()}


@app.get("/assistant/debug/retrieval")
def debug_retrieval(q: str = "agile team participants", top_k: int = 5):
    """
    Debug endpoint: test what the assistant retrieves for a query.
    GET /assistant/debug/retrieval?q=who+are+agile+team+members
    """
    try:
        # Try p2 path first
        try:
            from backend.services.p2.retrieval import search
        except ImportError:
            from backend.services.p2.retrieval import search
        results = search(query=q, top_k=top_k)
        return {
            "query":   q,
            "total":   len(results),
            "chunks":  [
                {
                    "doc_title":    r.get("doc_title"),
                    "section_name": r.get("section_name"),
                    "score":        r.get("score"),
                    "preview":      r.get("chunk_text", "")[:120],
                }
                for r in results
            ]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ══════════════════════════════════════════════════════════════════════════════
# ASSISTANT ROUTES  (mounted under /assistant prefix)
# ══════════════════════════════════════════════════════════════════════════════

asst = APIRouter(prefix="/assistant", tags=["Assistant"])


@asst.post("/chat")
def assistant_chat(payload: dict):
    """
    Send a message to the stateful assistant.

    payload:
      message    : str           — user message (required)
      thread_id  : str | null    — continue existing thread, or null to start new
      user_id    : str | null    — optional user identifier
      filters    : dict | null   — { doc_title, industry, doc_type }
      industry   : str | null    — industry context override

    Returns:
      thread_id, reply, sources, ticket_id, ticket_url, intent, messages, trace_id
    """
    try:
        message = payload.get("message", "").strip()
        if not message:
            raise HTTPException(status_code=400, detail="message is required")
        result = run_turn(
            user_message = message,
            thread_id    = payload.get("thread_id"),
            user_id      = payload.get("user_id"),
            filters      = payload.get("filters"),
            industry     = payload.get("industry"),
        )
        return result
    except HTTPException:
        raise
    except Exception as e:
        log.error(f"assistant_chat error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@asst.get("/thread/{thread_id}")
def get_thread(thread_id: str):
    """Return full conversation history for a thread."""
    from backend.services.p3.memory import restore_state
    state = restore_state(thread_id)
    if not state:
        raise HTTPException(status_code=404, detail="Thread not found")
    return {
        "thread_id": thread_id,
        "messages":  state.get("messages", []),
        "intent":    state.get("intent"),
        "industry":  state.get("industry"),
        "ticket_id": state.get("ticket_id"),
    }


@asst.delete("/thread/{thread_id}")
def clear_thread(thread_id: str):
    """Clear a conversation thread from Redis cache (keeps DB record)."""
    from backend.services.p3.memory import invalidate_state
    invalidate_state(thread_id)
    return {"message": f"Thread {thread_id} cleared from cache"}


@asst.get("/threads")
def list_assistant_threads(user_id: str = None, limit: int = 20):
    """List recent threads, optionally filtered by user_id."""
    threads = list_threads(user_id=user_id, limit=limit)
    return {"threads": threads, "total": len(threads)}


@asst.get("/tickets")
def get_tickets(limit: int = 20):
    """Fetch support tickets from Notion Ticket DB."""
    tickets = fetch_tickets(limit=limit)
    return {"tickets": tickets, "total": len(tickets)}


@asst.patch("/tickets/{ticket_id}/status")
def update_ticket(ticket_id: str, payload: dict):
    """Update ticket status: Open / In Progress / Resolved."""
    status = payload.get("status", "Open")
    if status not in ("Open", "In Progress", "Resolved"):
        raise HTTPException(status_code=400, detail="Invalid status")
    try:
        update_ticket_status(ticket_id, status)
        return {"message": f"Ticket {ticket_id} updated to '{status}'"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@asst.get("/log")
def get_assistant_log(limit: int = 50):
    """Fetch assistant conversation log from Notion Assistant Log DB."""
    try:
        entries = fetch_assistant_log(limit=limit)
        return {"entries": entries, "total": len(entries)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


app.include_router(asst)


# ══════════════════════════════════════════════════════════════════════════════
# CITERAG ROUTES  (mounted under /rag prefix)
# ══════════════════════════════════════════════════════════════════════════════

rag = APIRouter(prefix="/rag", tags=["CiteRAG"])


@rag.get("/health")
def rag_health():
    return {"status": "ok", "service": "CiteRAG"}


@rag.post("/ingest")
def rag_ingest(payload: dict = {}):
    from backend.services.p2.ingestion import ingest_notion
    try:
        page_ids = payload.get("page_ids") if payload else None
        force    = payload.get("force", False) if payload else False
        return ingest_notion(page_ids=page_ids, force=force)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@rag.get("/ingest/status")
def rag_ingest_status():
    try:
        return {"documents": list_ingested_docs()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@rag.delete("/ingest/{page_id}")
def rag_delete_ingested(page_id: str):
    try:
        count = delete_chunks_by_page(page_id)
        return {"deleted_chunks": count, "page_id": page_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@rag.post("/search")
def rag_search(payload: dict):
    from backend.services.p2.retrieval import search as do_search
    try:
        query   = payload["query"]
        top_k   = payload.get("top_k", 5)
        filters = payload.get("filters")
        cached  = redis_svc.get_cached_search(query, top_k, filters)
        if cached:
            return {"results": cached, "total": len(cached), "cached": True}
        results = do_search(query=query, top_k=top_k, filters=filters)
        redis_svc.cache_search(query, top_k, filters, results)
        return {"results": results, "total": len(results), "cached": False}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@rag.post("/ask")
def rag_ask(payload: dict):
    from backend.services.p2.qa import ask as do_ask
    from backend.services.p2.qa_log import log_qa
    try:
        question = payload["question"]
        top_k    = payload.get("top_k", 6)
        filters  = payload.get("filters")
        # Return cached answer if available (skips embedding + LLM call)
        cached = redis_svc.get_cached_ask(question, top_k, filters)
        if cached:
            cached["cached"] = True
            return cached
        result = do_ask(question=question, top_k=top_k, filters=filters)
        # Cache before logging so cache hit skips logging too
        redis_svc.cache_ask(question, top_k, filters, result)
        result["cached"] = False
        if result.get("sources"):
            try:
                log_entry = log_qa(
                    question = question,
                    answer   = result["answer"],
                    sources  = result["sources"],
                    filters  = filters,
                )
                result["notion_log_url"] = log_entry.get("url", "")
            except Exception as log_err:
                log.warning(f"Q&A log failed (non-fatal): {log_err}")
                result["notion_log_url"] = ""
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@rag.post("/ask/refine")
def rag_refine(payload: dict):
    from backend.services.p2.qa import refine_answer
    try:
        return refine_answer(
            question = payload["question"],
            answer   = payload["answer"],
            feedback = payload["feedback"],
            sources  = payload.get("sources", []),
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@rag.get("/qa-log")
def rag_qa_log(limit: int = 50):
    from backend.services.p2.qa_log import fetch_qa_log
    try:
        return {"entries": fetch_qa_log(limit=limit)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@rag.post("/compare")
def rag_compare(payload: dict):
    from backend.services.p2.compare import compare_documents
    from backend.services.p2.qa_log import log_compare
    try:
        title_a = payload["title_a"]
        title_b = payload["title_b"]
        focus   = payload.get("focus", "overall content and structure")
        cached  = redis_svc.get_cached_compare(title_a, title_b, focus)
        if cached:
            cached["cached"] = True
            return cached
        result = compare_documents(title_a=title_a, title_b=title_b, focus=focus)
        redis_svc.cache_compare(title_a, title_b, focus, result)
        result["cached"] = False
        try:
            log_compare(
                title_a    = title_a,
                title_b    = title_b,
                focus      = focus,
                comparison = result.get("comparison", ""),
                chunks_a   = result.get("doc_a_chunks", []),
                chunks_b   = result.get("doc_b_chunks", []),
            )
        except Exception as log_err:
            log.warning(f"log_compare failed (non-fatal): {log_err}")
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@rag.get("/documents")
def rag_documents():
    try:
        return {"titles": get_all_titles(), "industries": get_all_industries()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@rag.post("/eval/run")
def rag_eval_run(payload: dict):
    from backend.services.p2.eval import run_evaluation
    from backend.services.p2.qa_log import log_eval
    try:
        config = payload.get("config", {})
        result = run_evaluation(questions=payload["questions"], config=config)
        # Log each question individually to Q&A Log (non-blocking)
        run_name = config.get("run_name", "Eval run")
        try:
            for r in result.get("results", []):
                log_eval(
                    question = r.get("question", ""),
                    answer   = r.get("answer", ""),
                    sources  = r.get("sources", []),
                    scores   = {
                        k: r.get(k, 0)
                        for k in ("faithfulness", "answer_relevancy",
                                  "context_precision", "context_recall")
                        if k in r
                    },
                    run_name = run_name,
                )
        except Exception as log_err:
            log.warning(f"log_eval failed (non-fatal): {log_err}")
        return result
    except ImportError as e:
        raise HTTPException(status_code=422, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@rag.get("/eval/runs")
def rag_eval_runs():
    try:
        return {"runs": list_eval_runs()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@rag.get("/eval/runs/{run_id}")
def rag_eval_run_detail(run_id: int):
    try:
        run = get_eval_run(run_id)
        if not run:
            raise HTTPException(status_code=404, detail="Eval run not found")
        return run
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


app.include_router(rag)